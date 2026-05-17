"""
╔══════════════════════════════════════════════════════════════╗
║        CV ANALYZER & JOB APPLICATION SUITE                  ║
║        مجاني ١٠٠٪ — بدون اشتراك — للأبد                    ║
║        يعمل بـ Google Gemini أو Groq (كلاهم مجاني)          ║
╚══════════════════════════════════════════════════════════════╝

WHAT THIS APP DOES:
1. Analyzes CV vs job description → match score + gaps
2. Rewrites CV to 95%+ match
3. Generates Cover Letter tailored to company
4. Generates LinkedIn messages (3 variants)
5. Builds STAR interview stories
6. Builds Interview Q&A
7. Builds 30-day job action plan
8. Exports everything to DOCX files

REQUIREMENTS:
    pip install flask python-docx pillow requests

HOW TO RUN:
    python app.py
    Then open: http://localhost:5000
    Mobile (same WiFi): http://YOUR_IP:5000

FREE API KEYS (مجاناً):
  Gemini:  https://aistudio.google.com/app/apikey
  Groq:    https://console.groq.com  → API Keys
"""

import os, json, io, re
from flask import Flask, request, jsonify, send_file, render_template_string
import urllib.request, urllib.error
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# ─────────────────────────────────────────
# AI PROVIDER WRAPPERS  (Gemini + Groq)
# ─────────────────────────────────────────

GEMINI_MODELS = [
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-flash-8b",
]

GROQ_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-8b-instant",
    "mixtral-8x7b-32768",
]


def _http_post(url, headers, body_dict):
    data = json.dumps(body_dict).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8")
        raise RuntimeError(f"HTTP {e.code}: {err_body}")


def call_gemini(prompt, api_key, system=None, json_mode=False):
    """Call Google Gemini API (free tier)."""
    full_prompt = prompt
    if system:
        full_prompt = system + "\n\n" + prompt
    if json_mode:
        full_prompt += "\n\nCRITICAL: Respond ONLY with valid JSON. No markdown, no backticks, no preamble whatsoever."

    for model in GEMINI_MODELS:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        body = {
            "contents": [{"parts": [{"text": full_prompt}]}],
            "generationConfig": {
                "maxOutputTokens": 8192,
                "temperature": 0.7,
            }
        }
        try:
            result = _http_post(url, headers, body)
            text = result["candidates"][0]["content"]["parts"][0]["text"]
            return text.strip()
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                continue  # try next model
            raise
    raise RuntimeError("All Gemini models exhausted or quota exceeded.")


def call_groq(prompt, api_key, system=None, json_mode=False):
    """Call Groq API (free tier)."""
    if json_mode:
        prompt += "\n\nCRITICAL: Respond ONLY with valid JSON. No markdown, no backticks, no preamble."

    sys_msg = system or "You are an expert career coach and HR consultant specializing in UAE healthcare sector."
    messages = [
        {"role": "system", "content": sys_msg},
        {"role": "user", "content": prompt}
    ]

    for model in GROQ_MODELS:
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        body = {
            "model": model,
            "messages": messages,
            "max_tokens": 8192,
            "temperature": 0.7,
        }
        try:
            result = _http_post(url, headers, body)
            return result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            if "429" in str(e) or "rate" in str(e).lower() or "model" in str(e).lower():
                continue
            raise
    raise RuntimeError("All Groq models exhausted or rate-limited.")


def call_ai(prompt, provider, api_key, system=None, json_mode=False):
    """Unified AI caller — picks Gemini or Groq."""
    if not api_key:
        raise ValueError("لم يتم إدخال API Key")
    if provider == "gemini":
        raw = call_gemini(prompt, api_key, system=system, json_mode=json_mode)
    elif provider == "groq":
        raw = call_groq(prompt, api_key, system=system, json_mode=json_mode)
    else:
        raise ValueError(f"مزود غير معروف: {provider}")

    if json_mode:
        # Strip any accidental markdown fences
        raw = re.sub(r"^```(?:json)?\s*", "", raw.strip())
        raw = re.sub(r"\s*```$", "", raw.strip())
        return json.loads(raw)
    return raw


# ─────────────────────────────────────────
# TASK 1: ANALYZE CV vs JOB DESCRIPTION
# ─────────────────────────────────────────

def analyze_cv(cv_text, job_desc, provider, api_key):
    prompt = f"""
You are a senior HR consultant and ATS expert for UAE healthcare sector.

Analyze this CV against the job description below.

JOB DESCRIPTION:
{job_desc}

CV TEXT:
{cv_text}

Return a JSON object with EXACTLY this structure:
{{
  "match_score": <integer 0-100>,
  "match_label": "<Poor|Fair|Good|Strong|Excellent>",
  "strengths": [
    {{"title": "...", "detail": "..."}}
  ],
  "weaknesses": [
    {{"title": "...", "detail": "...", "impact": "High|Medium|Low"}}
  ],
  "missing_keywords": ["keyword1", "keyword2"],
  "suggestions": [
    {{"section": "Profile Summary|Job Title|Skills|Experience", "action": "...", "example": "..."}}
  ],
  "quick_wins": ["Quick change 1", "Quick change 2"],
  "overall_verdict": "One paragraph summary of candidacy fit"
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 2: REWRITE CV CONTENT
# ─────────────────────────────────────────

def rewrite_cv(cv_text, job_desc, candidate_name, provider, api_key):
    prompt = f"""
You are an expert CV writer for UAE healthcare executives.

Rewrite this CV to achieve 95%+ match with the job description.
Candidate: {candidate_name}

JOB DESCRIPTION:
{job_desc}

ORIGINAL CV:
{cv_text}

Rules:
- Keep ALL real facts, dates, companies — never invent
- Rewrite language to match JD keywords exactly
- Add commissioning, SOP development, referring physicians coordination if evidence exists
- Upgrade job titles strategically (e.g. "Clinic Manager" → "Senior Operations Manager")
- Add alignment statement to target company
- Every bullet must show a measurable result

Return JSON with this exact structure:
{{
  "header": {{
    "name": "...",
    "title": "...",
    "phone": "...",
    "email": "...",
    "location": "...",
    "education": "..."
  }},
  "summary": "Full rewritten summary paragraph...",
  "competencies": [
    {{"category": "...", "skills": "skill1 | skill2 | skill3"}}
  ],
  "career_highlights": {{
    "title": "...",
    "bullets": ["...", "..."]
  }},
  "experience": [
    {{
      "title": "...",
      "company": "...",
      "location": "...",
      "dates": "...",
      "reports_to": "...",
      "bullets": ["...", "..."]
    }}
  ],
  "education": [
    {{"degree": "...", "field": "...", "bullets": ["..."]}}
  ],
  "alignment_statement": "Company-specific paragraph about why you fit their mission..."
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 3: GENERATE COVER LETTER
# ─────────────────────────────────────────

def generate_cover_letter(cv_text, job_desc, candidate_name, company_name, role_title, provider, api_key):
    prompt = f"""
Write a powerful, tailored cover letter for:
- Candidate: {candidate_name}
- Role: {role_title}
- Company: {company_name}

JOB DESCRIPTION:
{job_desc}

CV SUMMARY:
{cv_text[:2000]}

Rules:
- Professional tone, 4-5 paragraphs
- Reference the company by name and their specific mission/values
- Include 1 value proposition table (6 bullet points matching JD requirements)
- Show knowledge of the role's specific challenges
- End with confident call to action
- Do NOT use generic phrases like "I am writing to apply"

Return JSON:
{{
  "subject": "RE: [Role Title] — [Candidate Name] Application",
  "opening": "Opening paragraph...",
  "value_props": [
    {{"check": "✓", "text": "Value proposition matching JD requirement"}}
  ],
  "body_paragraphs": ["paragraph 1...", "paragraph 2...", "paragraph 3..."],
  "closing": "Closing paragraph...",
  "signature": {{"name": "...", "title": "...", "phone": "...", "email": "...", "location": "..."}}
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 4: GENERATE LINKEDIN MESSAGES
# ─────────────────────────────────────────

def generate_linkedin_messages(candidate_name, role_title, company_name, top_achievement, provider, api_key):
    prompt = f"""
Generate 3 LinkedIn messages for {candidate_name} applying to {role_title} at {company_name}.
Top achievement to mention: {top_achievement}

Return JSON:
{{
  "headline": "Optimized LinkedIn headline (under 220 chars)",
  "about_section": "Full LinkedIn About/Summary section (5-7 paragraphs with bullets)",
  "messages": [
    {{
      "type": "Hiring Manager / Recruiter",
      "subject": "...",
      "body": "Full message text..."
    }},
    {{
      "type": "Connection Request Note (300 chars max)",
      "subject": null,
      "body": "Short connection note..."
    }},
    {{
      "type": "Internal Employee (Intel Gathering)",
      "subject": null,
      "body": "Message to current employee asking about culture..."
    }}
  ],
  "profile_tips": ["Tip 1...", "Tip 2...", "Tip 3..."]
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 5: GENERATE STAR STORIES
# ─────────────────────────────────────────

def generate_star_stories(cv_text, job_desc, provider, api_key):
    prompt = f"""
Based on this CV and job description, generate 5 STAR interview stories.
Each story should address a key competency from the JD.

JOB DESCRIPTION:
{job_desc}

CV:
{cv_text[:3000]}

Return JSON:
{{
  "stories": [
    {{
      "competency": "e.g., Commissioning & Facility Setup",
      "title": "Short story title",
      "situation": "...",
      "task": "...",
      "action": "...",
      "result": "Specific measurable outcomes...",
      "key_message": "One-line takeaway for interviewer"
    }}
  ]
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 6: GENERATE INTERVIEW Q&A
# ─────────────────────────────────────────

def generate_interview_qa(cv_text, job_desc, company_name, provider, api_key):
    prompt = f"""
Generate 6 likely interview questions for this role at {company_name}, with tailored answers.

JOB DESCRIPTION:
{job_desc}

CV CONTEXT:
{cv_text[:2000]}

Return JSON:
{{
  "questions": [
    {{
      "question": "...",
      "answer": "Full suggested answer (2-3 paragraphs)...",
      "tip": "Quick coaching tip for delivery"
    }}
  ],
  "questions_to_ask": [
    "Strategic question to ask interviewer 1",
    "Strategic question to ask interviewer 2",
    "Strategic question to ask interviewer 3"
  ]
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# TASK 7: GENERATE 30-DAY ACTION PLAN
# ─────────────────────────────────────────

def generate_action_plan(candidate_name, role_title, company_name, provider, api_key):
    prompt = f"""
Create a detailed 30-day action plan for {candidate_name} to secure the {role_title} role at {company_name}.

Return JSON:
{{
  "weeks": [
    {{
      "week": "Week 1 — Apply & Activate",
      "focus": "Brief focus description",
      "tasks": [
        {{"day": "Day 1-2", "action": "...", "platform": "LinkedIn|Email|Phone|Research", "priority": "HIGH|MEDIUM"}}
      ]
    }}
  ],
  "pro_tips": ["Pro tip 1...", "Pro tip 2..."],
  "red_flags_to_avoid": ["Don't do this...", "Avoid this mistake..."],
  "success_metrics": ["How to know you're on track..."]
}}
"""
    return call_ai(prompt, provider, api_key, json_mode=True)


# ─────────────────────────────────────────
# DOCX GENERATORS
# ─────────────────────────────────────────

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading_line(doc, text, level=1, color_hex="1B4F8C"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    r, g, b = hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def generate_cv_docx(cv_data):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    h = cv_data.get("header", {})
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_r = name_p.add_run(h.get("name", "").upper())
    name_r.bold = True
    name_r.font.size = Pt(22)
    name_r.font.color.rgb = RGBColor(27, 79, 140)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run(h.get("title", ""))
    title_r.bold = True
    title_r.font.size = Pt(13)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_r = contact_p.add_run(f"{h.get('phone','')}  |  {h.get('email','')}  |  {h.get('location','')}")
    contact_r.font.size = Pt(10)
    contact_r.font.color.rgb = RGBColor(85, 85, 85)

    edu_p = doc.add_paragraph()
    edu_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edu_r = edu_p.add_run(h.get("education", ""))
    edu_r.font.size = Pt(10)
    edu_r.font.color.rgb = RGBColor(85, 85, 85)

    add_heading_line(doc, "PROFESSIONAL SUMMARY")
    s = doc.add_paragraph(cv_data.get("summary", ""))
    s.paragraph_format.space_after = Pt(6)
    for run in s.runs:
        run.font.size = Pt(10)

    if cv_data.get("competencies"):
        add_heading_line(doc, "CORE COMPETENCIES")
        table = doc.add_table(rows=len(cv_data["competencies"]), cols=2)
        table.style = 'Table Grid'
        for i, comp in enumerate(cv_data["competencies"]):
            row = table.rows[i]
            c0 = row.cells[0]
            c1 = row.cells[1]
            set_cell_bg(c0, "D6E4F0")
            c0.width = Inches(1.8)
            c1.width = Inches(5.2)
            p0 = c0.paragraphs[0]
            r0 = p0.add_run(comp.get("category", ""))
            r0.bold = True
            r0.font.size = Pt(9)
            r0.font.color.rgb = RGBColor(27, 79, 140)
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(comp.get("skills", ""))
            r1.font.size = Pt(9)

    if cv_data.get("career_highlights"):
        add_heading_line(doc, "CAREER HIGHLIGHTS")
        ch = cv_data["career_highlights"]
        ht = doc.add_paragraph()
        ht_r = ht.add_run(ch.get("title", ""))
        ht_r.bold = True
        ht_r.font.size = Pt(10)
        for b in ch.get("bullets", []):
            add_bullet(doc, b)

    if cv_data.get("experience"):
        add_heading_line(doc, "PROFESSIONAL EXPERIENCE")
        for exp in cv_data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r1 = p.add_run(exp.get("title", "") + " — ")
            r1.bold = True
            r1.font.color.rgb = RGBColor(27, 79, 140)
            r1.font.size = Pt(11)
            r2 = p.add_run(exp.get("company", ""))
            r2.bold = True
            r2.font.size = Pt(11)
            r3 = p.add_run("  |  " + exp.get("location", ""))
            r3.font.size = Pt(10)
            r3.font.color.rgb = RGBColor(85, 85, 85)
            dp = doc.add_paragraph()
            dr = dp.add_run(exp.get("dates", "") + "  |  Reporting to: " + exp.get("reports_to", ""))
            dr.italic = True
            dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(85, 85, 85)
            dp.paragraph_format.space_after = Pt(4)
            for b in exp.get("bullets", []):
                add_bullet(doc, b)

    if cv_data.get("education"):
        add_heading_line(doc, "EDUCATION & PROFESSIONAL DEVELOPMENT")
        for edu in cv_data["education"]:
            p = doc.add_paragraph()
            r = p.add_run(edu.get("degree", "") + " — " + edu.get("field", ""))
            r.bold = True
            r.font.size = Pt(10)
            for b in edu.get("bullets", []):
                add_bullet(doc, b)

    if cv_data.get("alignment_statement"):
        add_heading_line(doc, "ALIGNMENT WITH TARGET ORGANIZATION")
        ap = doc.add_paragraph(cv_data["alignment_statement"])
        for run in ap.runs:
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_cover_letter_docx(cl_data, candidate_info):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(candidate_info.get("name", "").upper())
    nr.bold = True
    nr.font.size = Pt(20)
    nr.font.color.rgb = RGBColor(27, 79, 140)

    sig = cl_data.get("signature", {})
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f"{sig.get('phone','')}  |  {sig.get('email','')}  |  {sig.get('location','')}")
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(85, 85, 85)

    dp = doc.add_paragraph()
    pPr = dp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '1B4F8C')
    pBdr.append(bot); pPr.append(pBdr)

    doc.add_paragraph()
    sp = doc.add_paragraph()
    sr = sp.add_run(cl_data.get("subject", ""))
    sr.bold = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(27, 79, 140)

    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")

    op = doc.add_paragraph(cl_data.get("opening", ""))
    op.paragraph_format.space_after = Pt(10)

    vps = cl_data.get("value_props", [])
    if vps:
        vp_label = doc.add_paragraph()
        vlr = vp_label.add_run("What I Bring to This Role:")
        vlr.bold = True
        vlr.font.size = Pt(11)
        vlr.font.color.rgb = RGBColor(27, 79, 140)
        rows = [vps[i:i+2] for i in range(0, len(vps), 2)]
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = 'Table Grid'
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, vp in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    if i % 2 == 0:
                        set_cell_bg(cell, "D6E4F0")
                    p = cell.paragraphs[0]
                    r = p.add_run(vp.get("check", "✓") + "  " + vp.get("text", ""))
                    r.font.size = Pt(9)

    doc.add_paragraph()
    for para_text in cl_data.get("body_paragraphs", []):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(10)

    cp2 = doc.add_paragraph(cl_data.get("closing", ""))
    for run in cp2.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()

    np = doc.add_paragraph()
    nr2 = np.add_run(sig.get("name", ""))
    nr2.bold = True
    nr2.font.size = Pt(12)
    nr2.font.color.rgb = RGBColor(27, 79, 140)

    for field in ["title", "phone", "email", "location"]:
        fp = doc.add_paragraph(sig.get(field, ""))
        for run in fp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(85, 85, 85)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_strategy_docx(linkedin_data, star_data, qa_data, plan_data):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("JOB APPLICATION STRATEGY PACK")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor(27, 79, 140)

    add_heading_line(doc, "SECTION 1 — LinkedIn Profile Optimization")
    if linkedin_data.get("headline"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "D6E4F0")
        tbl.rows[0].cells[0].paragraphs[0].add_run(linkedin_data["headline"]).font.size = Pt(10)

    if linkedin_data.get("about_section"):
        doc.add_paragraph()
        about_h = doc.add_paragraph()
        about_h.add_run("Recommended About Section:").bold = True
        tbl2 = doc.add_table(rows=1, cols=1)
        tbl2.style = 'Table Grid'
        set_cell_bg(tbl2.rows[0].cells[0], "F5F5F5")
        tbl2.rows[0].cells[0].paragraphs[0].add_run(linkedin_data["about_section"]).font.size = Pt(9)

    for tip in linkedin_data.get("profile_tips", []):
        add_bullet(doc, tip)

    add_heading_line(doc, "SECTION 2 — LinkedIn Outreach Messages")
    for msg in linkedin_data.get("messages", []):
        doc.add_paragraph()
        mh = doc.add_paragraph()
        mh.add_run(f"Message: {msg.get('type','')}").bold = True
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "EBF5FB")
        body_text = (f"Subject: {msg['subject']}\n\n" if msg.get('subject') else "") + msg.get("body", "")
        tbl.rows[0].cells[0].paragraphs[0].add_run(body_text).font.size = Pt(9)

    add_heading_line(doc, "SECTION 3 — STAR Interview Stories")
    for story in star_data.get("stories", []):
        doc.add_paragraph()
        sh = doc.add_paragraph()
        sr = sh.add_run(f"Competency: {story.get('competency','')} — {story.get('title','')}")
        sr.bold = True
        sr.font.color.rgb = RGBColor(27, 79, 140)
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = 'Table Grid'
        labels = ["SITUATION", "TASK", "ACTION", "RESULT"]
        keys = ["situation", "task", "action", "result"]
        colors = ["D6E4F0", "F5F5F5", "D6E4F0", "D6F0E0"]
        for i, (lbl, key, col) in enumerate(zip(labels, keys, colors)):
            row = tbl.rows[i]
            set_cell_bg(row.cells[0], col)
            row.cells[0].width = Inches(1.2)
            row.cells[0].paragraphs[0].add_run(lbl).bold = True
            row.cells[1].paragraphs[0].add_run(story.get(key, "")).font.size = Pt(9)
        km = doc.add_paragraph()
        kmr = km.add_run("Key Message: " + story.get("key_message", ""))
        kmr.italic = True
        kmr.font.size = Pt(9)
        kmr.font.color.rgb = RGBColor(27, 79, 140)

    add_heading_line(doc, "SECTION 4 — Interview Questions & Answers")
    for qa in qa_data.get("questions", []):
        doc.add_paragraph()
        qp = doc.add_paragraph()
        qr = qp.add_run("Q: " + qa.get("question", ""))
        qr.bold = True
        qr.font.color.rgb = RGBColor(27, 79, 140)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "F0F8FF")
        tbl.rows[0].cells[0].paragraphs[0].add_run(qa.get("answer", "")).font.size = Pt(9)
        tip_p = doc.add_paragraph()
        tip_r = tip_p.add_run("💡 " + qa.get("tip", ""))
        tip_r.italic = True
        tip_r.font.size = Pt(9)
        tip_r.font.color.rgb = RGBColor(85, 85, 85)

    if qa_data.get("questions_to_ask"):
        add_heading_line(doc, "Questions to Ask the Interviewer", level=2)
        for q in qa_data["questions_to_ask"]:
            add_bullet(doc, q)

    add_heading_line(doc, "SECTION 5 — 30-Day Action Plan")
    for week in plan_data.get("weeks", []):
        doc.add_paragraph()
        wh = doc.add_paragraph()
        wr = wh.add_run(week.get("week", ""))
        wr.bold = True
        wr.font.color.rgb = RGBColor(27, 79, 140)
        fp = doc.add_paragraph(week.get("focus", ""))
        if fp.runs:
            fp.runs[0].italic = True
        for task in week.get("tasks", []):
            p = doc.add_paragraph(style='List Bullet')
            r1 = p.add_run(f"[{task.get('day','')}] ")
            r1.bold = True
            r1.font.size = Pt(9)
            r2 = p.add_run(task.get("action", ""))
            r2.font.size = Pt(9)
            if task.get("platform"):
                r3 = p.add_run(f" [{task['platform']}]")
                r3.font.size = Pt(8)
                r3.font.color.rgb = RGBColor(85, 85, 85)

    if plan_data.get("pro_tips"):
        add_heading_line(doc, "Pro Tips", level=2)
        for tip in plan_data["pro_tips"]:
            add_bullet(doc, "💡 " + tip)

    if plan_data.get("red_flags_to_avoid"):
        add_heading_line(doc, "Common Mistakes to Avoid", level=2)
        for flag in plan_data["red_flags_to_avoid"]:
            add_bullet(doc, "⚠️ " + flag)

    doc.add_paragraph()
    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp2.add_run("You are not just a candidate. You are the solution they are looking for.")
    fr.bold = True
    fr.font.size = Pt(13)
    fr.font.color.rgb = RGBColor(27, 79, 140)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    try:
        data = request.json
        provider = data.get('provider', 'gemini')
        api_key = data.get('api_key', '')
        cv_text = data.get('cv_text', '')
        job_desc = data.get('job_desc', '')
        result = analyze_cv(cv_text, job_desc, provider, api_key)
        return jsonify({"success": True, "data": result})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/generate_all', methods=['POST'])
def api_generate_all():
    try:
        data = request.json
        provider = data.get('provider', 'gemini')
        api_key = data.get('api_key', '')
        cv_text = data.get('cv_text', '')
        job_desc = data.get('job_desc', '')
        candidate_name = data.get('candidate_name', 'Candidate')
        company_name = data.get('company_name', 'Target Company')
        role_title = data.get('role_title', 'Target Role')
        top_achievement = data.get('top_achievement', '')

        results = {}
        results['cv'] = rewrite_cv(cv_text, job_desc, candidate_name, provider, api_key)
        results['cover_letter'] = generate_cover_letter(
            cv_text, job_desc, candidate_name, company_name, role_title, provider, api_key)
        results['linkedin'] = generate_linkedin_messages(
            candidate_name, role_title, company_name, top_achievement, provider, api_key)
        results['star'] = generate_star_stories(cv_text, job_desc, provider, api_key)
        results['qa'] = generate_interview_qa(cv_text, job_desc, company_name, provider, api_key)
        results['plan'] = generate_action_plan(candidate_name, role_title, company_name, provider, api_key)

        return jsonify({"success": True, "data": results})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/download/<doc_type>', methods=['POST'])
def api_download(doc_type):
    try:
        data = request.json
        if doc_type == 'cv':
            buf = generate_cv_docx(data['cv_data'])
            filename = "CV_Optimized.docx"
        elif doc_type == 'cover_letter':
            buf = generate_cover_letter_docx(data['cl_data'], data.get('candidate_info', {}))
            filename = "Cover_Letter.docx"
        elif doc_type == 'strategy':
            buf = generate_strategy_docx(
                data['linkedin_data'], data['star_data'],
                data['qa_data'], data['plan_data'])
            filename = "Job_Strategy_Pack.docx"
        else:
            return jsonify({"error": "Unknown doc type"}), 400

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─────────────────────────────────────────
# HTML FRONTEND
# ─────────────────────────────────────────

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0">
<title>CV Analyzer Pro — مجاني ١٠٠٪</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800&family=Space+Mono:wght@400;700&display=swap');

  :root {
    --blue: #1B4F8C;
    --blue-light: #D6E4F0;
    --blue-mid: #2E6DB4;
    --green: #1A6B3C;
    --green-light: #D6F0E0;
    --orange: #C0550A;
    --red: #B91C1C;
    --gold: #B8860B;
    --dark: #0F1923;
    --gray: #6B7280;
    --light: #F8FAFC;
    --white: #FFFFFF;
    --border: #E2E8F0;
    --shadow: 0 4px 24px rgba(27,79,140,0.12);
    --radius: 16px;
    --gemini: #4285F4;
    --groq: #FF6B35;
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }

  body {
    font-family: 'Tajawal', sans-serif;
    background: linear-gradient(135deg, #0F1923 0%, #1B4F8C 50%, #0F1923 100%);
    min-height: 100vh;
    color: var(--dark);
  }

  .app-header {
    background: rgba(255,255,255,0.05);
    backdrop-filter: blur(20px);
    border-bottom: 1px solid rgba(255,255,255,0.1);
    padding: 16px 20px;
    display: flex;
    align-items: center;
    gap: 12px;
    position: sticky;
    top: 0;
    z-index: 100;
  }
  .app-logo {
    width: 44px; height: 44px;
    background: var(--blue);
    border-radius: 12px;
    display: flex; align-items: center; justify-content: center;
    font-size: 22px; flex-shrink: 0;
  }
  .app-title { color: white; flex: 1; }
  .app-title h1 { font-size: 18px; font-weight: 800; letter-spacing: -0.3px; }
  .app-title p { font-size: 11px; opacity: 0.6; margin-top: 1px; }
  .free-badge {
    background: linear-gradient(135deg, #22c55e, #16a34a);
    color: white; font-size: 11px; font-weight: 800;
    padding: 4px 10px; border-radius: 20px;
    letter-spacing: 0.5px;
  }

  .main { max-width: 900px; margin: 0 auto; padding: 20px 16px 40px; }

  .card {
    background: white;
    border-radius: var(--radius);
    padding: 24px;
    margin-bottom: 16px;
    box-shadow: var(--shadow);
  }
  .card-header {
    display: flex; align-items: center; gap: 10px;
    margin-bottom: 18px;
  }
  .card-icon {
    width: 38px; height: 38px;
    background: var(--blue-light);
    border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 18px; flex-shrink: 0;
  }
  .card-title { font-size: 16px; font-weight: 700; color: var(--blue); }
  .card-subtitle { font-size: 12px; color: var(--gray); }

  label { font-size: 13px; font-weight: 600; color: var(--dark); margin-bottom: 6px; display: block; }
  input[type=text], input[type=password], textarea, select {
    width: 100%; padding: 12px 14px;
    border: 2px solid var(--border);
    border-radius: 10px;
    font-family: 'Tajawal', sans-serif;
    font-size: 14px; color: var(--dark);
    transition: border-color 0.2s;
    margin-bottom: 14px;
    background: var(--light);
    direction: ltr;
  }
  input:focus, textarea:focus { outline: none; border-color: var(--blue); background: white; }
  textarea { min-height: 120px; resize: vertical; }
  .form-row { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
  @media(max-width:600px) { .form-row { grid-template-columns: 1fr; } }

  /* ── PROVIDER SELECTOR ── */
  .provider-section {
    background: linear-gradient(135deg, #0F1923, #1B3A6B);
    border-radius: 16px; padding: 20px;
    margin-bottom: 16px; color: white;
  }
  .provider-section h3 {
    font-size: 14px; font-weight: 800; margin-bottom: 4px;
    color: white;
  }
  .provider-section p {
    font-size: 12px; opacity: 0.7; margin-bottom: 16px;
  }
  .provider-tabs {
    display: flex; gap: 8px; margin-bottom: 16px;
  }
  .provider-tab {
    flex: 1; padding: 12px 8px;
    border-radius: 12px; border: 2px solid rgba(255,255,255,0.15);
    background: rgba(255,255,255,0.05);
    color: rgba(255,255,255,0.6);
    cursor: pointer; text-align: center;
    font-family: 'Tajawal', sans-serif;
    font-size: 13px; font-weight: 700;
    transition: all 0.2s;
  }
  .provider-tab.active-gemini {
    border-color: var(--gemini);
    background: rgba(66,133,244,0.15);
    color: #93C5FD;
  }
  .provider-tab.active-groq {
    border-color: var(--groq);
    background: rgba(255,107,53,0.15);
    color: #FCA5A5;
  }
  .provider-tab:hover { border-color: rgba(255,255,255,0.3); }
  .provider-logo { font-size: 20px; display: block; margin-bottom: 4px; }
  .provider-free { font-size: 10px; opacity: 0.8; font-weight: 400; }

  .api-key-box { position: relative; }
  .api-key-box input {
    background: rgba(255,255,255,0.08);
    border-color: rgba(255,255,255,0.15);
    color: white; margin-bottom: 6px;
  }
  .api-key-box input::placeholder { color: rgba(255,255,255,0.35); }
  .api-key-box input:focus { background: rgba(255,255,255,0.12); border-color: rgba(255,255,255,0.4); }
  .api-link { color: #60A5FA; font-size: 12px; text-decoration: underline; cursor: pointer; }
  .get-key-hint {
    font-size: 12px; color: rgba(255,255,255,0.55);
    margin-top: -4px; margin-bottom: 0;
  }

  .btn {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 14px 24px; border-radius: 12px;
    font-size: 14px; font-weight: 700;
    font-family: 'Tajawal', sans-serif;
    cursor: pointer; border: none;
    transition: all 0.2s; text-decoration: none;
  }
  .btn-primary {
    background: var(--blue); color: white;
    width: 100%; justify-content: center;
    font-size: 15px; padding: 16px;
  }
  .btn-primary:hover { background: var(--blue-mid); transform: translateY(-1px); }
  .btn-primary:disabled { opacity: 0.5; cursor: not-allowed; transform: none; }
  .btn-outline {
    background: white; color: var(--blue);
    border: 2px solid var(--blue);
    width: 100%; justify-content: center; margin-bottom: 10px;
  }
  .btn-outline:hover { background: var(--blue-light); }

  .tabs { display: flex; gap: 6px; margin-bottom: 20px; flex-wrap: wrap; }
  .tab {
    padding: 8px 16px; border-radius: 20px;
    font-size: 13px; font-weight: 600;
    cursor: pointer; border: 2px solid var(--border);
    background: white; color: var(--gray);
    transition: all 0.2s;
    font-family: 'Tajawal', sans-serif;
  }
  .tab.active { background: var(--blue); color: white; border-color: var(--blue); }

  .score-circle {
    width: 120px; height: 120px; border-radius: 50%;
    display: flex; flex-direction: column;
    align-items: center; justify-content: center;
    margin: 0 auto 16px; position: relative;
  }
  .score-num { font-size: 42px; font-weight: 800; line-height: 1; font-family: 'Space Mono', monospace; }
  .score-pct { font-size: 14px; font-weight: 600; }
  .score-excellent { background: var(--green-light); color: var(--green); border: 4px solid var(--green); }
  .score-good { background: #FEF9C3; color: #854D0E; border: 4px solid #CA8A04; }
  .score-fair { background: #FEE2E2; color: var(--red); border: 4px solid var(--red); }

  .result-list { list-style: none; }
  .result-list li {
    padding: 12px 14px; border-radius: 10px;
    margin-bottom: 8px; font-size: 13px;
    line-height: 1.5; direction: ltr;
  }
  .list-strength { background: var(--green-light); border-right: 4px solid var(--green); }
  .list-weakness { background: #FEE2E2; border-right: 4px solid var(--red); }
  .list-suggestion { background: var(--blue-light); border-right: 4px solid var(--blue); }
  .list-item-title { font-weight: 700; margin-bottom: 3px; }
  .badge { display: inline-block; padding: 2px 8px; border-radius: 20px; font-size: 10px; font-weight: 700; margin-right: 6px; }
  .badge-high { background: #FEE2E2; color: var(--red); }
  .badge-medium { background: #FEF9C3; color: #854D0E; }
  .badge-low { background: var(--green-light); color: var(--green); }

  .keywords { display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; }
  .keyword { padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; background: #FEE2E2; color: var(--red); border: 1px solid #FECACA; }

  .loader { display: none; text-align: center; padding: 40px 20px; }
  .loader.show { display: block; }
  .spinner { width: 50px; height: 50px; border: 4px solid var(--blue-light); border-top-color: var(--blue); border-radius: 50%; animation: spin 0.8s linear infinite; margin: 0 auto 16px; }
  @keyframes spin { to { transform: rotate(360deg); } }
  .loader-text { font-size: 14px; color: var(--gray); }
  .loader-step { font-size: 12px; color: var(--blue); margin-top: 6px; font-weight: 600; }

  .tab-content { display: none; }
  .tab-content.active { display: block; }

  .output-section { border: 1px solid var(--border); border-radius: 12px; overflow: hidden; margin-bottom: 16px; }
  .output-section-header { background: var(--blue); color: white; padding: 12px 16px; font-size: 14px; font-weight: 700; display: flex; align-items: center; gap: 8px; cursor: pointer; }
  .output-section-body { padding: 16px; font-size: 13px; line-height: 1.7; direction: ltr; white-space: pre-wrap; background: var(--light); display: none; }
  .output-section-body.open { display: block; }

  .download-bar { background: linear-gradient(135deg, var(--blue), var(--blue-mid)); border-radius: 16px; padding: 20px; color: white; margin-bottom: 16px; }
  .download-bar h3 { font-size: 16px; font-weight: 800; margin-bottom: 14px; }
  .download-btns { display: flex; gap: 10px; flex-wrap: wrap; }
  .dl-btn { background: white; color: var(--blue); border: none; border-radius: 10px; padding: 10px 18px; font-size: 13px; font-weight: 700; cursor: pointer; display: flex; align-items: center; gap: 6px; font-family: 'Tajawal', sans-serif; transition: all 0.2s; }
  .dl-btn:hover { transform: translateY(-2px); box-shadow: 0 4px 12px rgba(0,0,0,0.15); }

  .alert { padding: 12px 16px; border-radius: 10px; font-size: 13px; margin-bottom: 14px; display: none; }
  .alert.show { display: block; }
  .alert-error { background: #FEE2E2; color: var(--red); border: 1px solid #FECACA; }
  .alert-success { background: var(--green-light); color: var(--green); border: 1px solid #BBF7D0; }

  .star-card { border: 1px solid var(--border); border-radius: 12px; overflow: hidden; margin-bottom: 16px; }
  .star-card-header { background: var(--blue); color: white; padding: 12px 16px; font-size: 14px; font-weight: 700; }
  .star-row { display: grid; grid-template-columns: 100px 1fr; border-bottom: 1px solid var(--border); }
  .star-row:last-child { border-bottom: none; }
  .star-label { background: var(--blue-light); color: var(--blue); font-weight: 700; font-size: 12px; padding: 12px; display: flex; align-items: center; }
  .star-text { padding: 12px; font-size: 13px; line-height: 1.6; direction: ltr; }
  .star-result { background: var(--green-light); }

  .progress-steps { display: flex; gap: 8px; margin-bottom: 20px; overflow-x: auto; padding-bottom: 4px; }
  .progress-step { display: flex; flex-direction: column; align-items: center; min-width: 70px; }
  .progress-dot { width: 32px; height: 32px; border-radius: 50%; background: var(--border); color: var(--gray); display: flex; align-items: center; justify-content: center; font-size: 14px; font-weight: 700; margin-bottom: 4px; transition: all 0.3s; }
  .progress-dot.done { background: var(--green); color: white; }
  .progress-dot.active { background: var(--blue); color: white; animation: pulse 1s infinite; }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:0.7} }
  .progress-label { font-size: 9px; color: var(--gray); text-align: center; }

  @media(max-width: 480px) {
    .card { padding: 16px; }
    .score-circle { width: 100px; height: 100px; }
    .score-num { font-size: 36px; }
  }
</style>
</head>
<body>

<div class="app-header">
  <div class="app-logo">🎯</div>
  <div class="app-title">
    <h1>CV Analyzer Pro</h1>
    <p>محلل السيرة الذاتية الذكي</p>
  </div>
  <span class="free-badge">FREE ∞</span>
</div>

<div class="main">

  <!-- ══ PROVIDER SELECTOR ══ -->
  <div class="provider-section">
    <h3>🤖 اختر مزود الذكاء الاصطناعي</h3>
    <p>كلاهم مجاني تماماً — لا بطاقة ائتمان — لا اشتراك</p>

    <div class="provider-tabs">
      <div class="provider-tab active-gemini" id="tabGemini" onclick="selectProvider('gemini')">
        <span class="provider-logo">🔵</span>
        Google Gemini
        <div class="provider-free">مجاني — 15 req/min</div>
      </div>
      <div class="provider-tab" id="tabGroq" onclick="selectProvider('groq')">
        <span class="provider-logo">🟠</span>
        Groq
        <div class="provider-free">مجاني — سريع جداً</div>
      </div>
    </div>

    <div class="api-key-box">
      <label id="apiKeyLabel" style="color:rgba(255,255,255,0.8)">🔑 Gemini API Key</label>
      <input type="password" id="apiKey" placeholder="AIza..." autocomplete="off">
      <p class="get-key-hint" id="apiKeyHint">
        احصل على مفتاح مجاني: 
        <a class="api-link" id="apiKeyLink" href="https://aistudio.google.com/app/apikey" target="_blank">
          aistudio.google.com
        </a>
        — سجّل بحساب Google وخذ المفتاح فوراً
      </p>
    </div>
  </div>

  <!-- ══ STEP 1: INPUT ══ -->
  <div id="stepInput" class="card">
    <div class="card-header">
      <div class="card-icon">📋</div>
      <div>
        <div class="card-title">الخطوة ١ — أدخل البيانات</div>
        <div class="card-subtitle">السيرة الذاتية + وصف الوظيفة</div>
      </div>
    </div>

    <div class="form-row">
      <div>
        <label>اسمك الكامل</label>
        <input type="text" id="candidateName" placeholder="Haitham El Meslemani">
      </div>
      <div>
        <label>اسم الشركة المستهدفة</label>
        <input type="text" id="companyName" placeholder="M42 / HealthPlus">
      </div>
    </div>

    <div class="form-row">
      <div>
        <label>المسمى الوظيفي المستهدف</label>
        <input type="text" id="roleTitle" placeholder="Senior Operations Manager">
      </div>
      <div>
        <label>أبرز إنجاز واحد</label>
        <input type="text" id="topAchievement" placeholder="Scaled clinic from 1.2M to 55M AED">
      </div>
    </div>

    <label>نص السيرة الذاتية (الصق هنا)</label>
    <textarea id="cvText" placeholder="الصق نص سيرتك الذاتية هنا..."></textarea>

    <label>وصف الوظيفة Job Description</label>
    <textarea id="jobDesc" placeholder="الصق وصف الوظيفة هنا..."></textarea>

    <div id="errorAlert" class="alert alert-error">⚠️ خطأ: <span id="errorMsg"></span></div>

    <button class="btn btn-outline" onclick="runAnalysis()">
      🔍 تحليل التطابق فقط (سريع)
    </button>

    <button class="btn btn-primary" onclick="runFullGeneration()">
      ⚡ توليد كل الوثائق دفعة واحدة
    </button>
  </div>

  <!-- LOADER -->
  <div id="loader" class="card loader">
    <div class="spinner"></div>
    <div class="loader-text" id="loaderText">جاري العمل...</div>
    <div class="loader-step" id="loaderStep">يرجى الانتظار</div>
    <div class="progress-steps" id="progressSteps" style="margin-top:20px;">
      <div class="progress-step"><div class="progress-dot" id="dot1">1</div><div class="progress-label">تحليل السيرة</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot2">2</div><div class="progress-label">إعادة الكتابة</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot3">3</div><div class="progress-label">Cover Letter</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot4">4</div><div class="progress-label">LinkedIn</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot5">5</div><div class="progress-label">STAR Stories</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot6">6</div><div class="progress-label">مقابلة Q&A</div></div>
      <div class="progress-step"><div class="progress-dot" id="dot7">7</div><div class="progress-label">خطة ٣٠ يوم</div></div>
    </div>
  </div>

  <!-- RESULTS -->
  <div id="results" style="display:none;">

    <div id="analysisCard" class="card" style="display:none;">
      <div class="card-header">
        <div class="card-icon">📊</div>
        <div>
          <div class="card-title">نتيجة التحليل</div>
          <div class="card-subtitle">CV vs Job Description</div>
        </div>
      </div>
      <div id="scoreDisplay"></div>
      <div id="analysisDetails"></div>
    </div>

    <div id="downloadBar" class="download-bar" style="display:none;">
      <h3>⬇️ تحميل الوثائق</h3>
      <div class="download-btns">
        <button class="dl-btn" onclick="downloadDoc('cv')">📄 CV المُحسَّنة</button>
        <button class="dl-btn" onclick="downloadDoc('cover_letter')">✉️ Cover Letter</button>
        <button class="dl-btn" onclick="downloadDoc('strategy')">🎯 حزمة الاستراتيجية</button>
      </div>
    </div>

    <div id="outputTabs" class="card" style="display:none;">
      <div class="tabs">
        <div class="tab active" onclick="switchTab('tabCV', this)">📄 السيرة</div>
        <div class="tab" onclick="switchTab('tabCL', this)">✉️ Cover Letter</div>
        <div class="tab" onclick="switchTab('tabLI', this)">💼 LinkedIn</div>
        <div class="tab" onclick="switchTab('tabSTAR', this)">⭐ STAR</div>
        <div class="tab" onclick="switchTab('tabQA', this)">🎤 مقابلة</div>
        <div class="tab" onclick="switchTab('tabPlan', this)">📅 خطة ٣٠ يوم</div>
      </div>

      <div id="tabCV" class="tab-content active"><div id="cvOutput"></div></div>
      <div id="tabCL" class="tab-content"><div id="clOutput"></div></div>
      <div id="tabLI" class="tab-content"><div id="liOutput"></div></div>
      <div id="tabSTAR" class="tab-content"><div id="starOutput"></div></div>
      <div id="tabQA" class="tab-content"><div id="qaOutput"></div></div>
      <div id="tabPlan" class="tab-content"><div id="planOutput"></div></div>
    </div>

  </div>
</div>

<script>
// ── STATE ──
let selectedProvider = 'gemini';
let generatedData = {};

// ── PROVIDER SWITCH ──
function selectProvider(p) {
  selectedProvider = p;
  document.getElementById('tabGemini').className = 'provider-tab' + (p==='gemini' ? ' active-gemini' : '');
  document.getElementById('tabGroq').className   = 'provider-tab' + (p==='groq'   ? ' active-groq'   : '');
  
  const label = document.getElementById('apiKeyLabel');
  const hint  = document.getElementById('apiKeyHint');
  const link  = document.getElementById('apiKeyLink');
  const input = document.getElementById('apiKey');
  
  if (p === 'gemini') {
    label.textContent = '🔑 Gemini API Key';
    input.placeholder = 'AIza...';
    link.href = 'https://aistudio.google.com/app/apikey';
    link.textContent = 'aistudio.google.com';
    hint.innerHTML = 'احصل على مفتاح مجاني: <a class="api-link" href="https://aistudio.google.com/app/apikey" target="_blank">aistudio.google.com</a> — سجّل بحساب Google وخذ المفتاح فوراً';
  } else {
    label.textContent = '🔑 Groq API Key';
    input.placeholder = 'gsk_...';
    hint.innerHTML = 'احصل على مفتاح مجاني: <a class="api-link" href="https://console.groq.com" target="_blank">console.groq.com</a> → API Keys → Create Key';
  }
  input.value = '';
}

// ── HELPERS ──
function showError(msg) {
  const el = document.getElementById('errorAlert');
  document.getElementById('errorMsg').textContent = msg;
  el.classList.add('show');
}
function hideError() {
  document.getElementById('errorAlert').classList.remove('show');
}
function setLoader(show, text='', step='') {
  const el = document.getElementById('loader');
  el.classList.toggle('show', show);
  if (text) document.getElementById('loaderText').textContent = text;
  if (step) document.getElementById('loaderStep').textContent = step;
}
function setDot(n, state) {
  const el = document.getElementById('dot'+n);
  if (!el) return;
  el.className = 'progress-dot' + (state==='done' ? ' done' : state==='active' ? ' active' : '');
  if (state === 'done') el.textContent = '✓';
}
function collapsibleSection(title, content, icon='') {
  const id = 'sec_' + Math.random().toString(36).slice(2,7);
  return `
    <div class="output-section">
      <div class="output-section-header" onclick="toggleSection('${id}')">
        ${icon ? icon+' ' : ''}${title} <span style="margin-right:auto;font-size:12px;opacity:0.7">اضغط للفتح ▼</span>
      </div>
      <div id="${id}" class="output-section-body">${content}</div>
    </div>`;
}
function toggleSection(id) {
  document.getElementById(id).classList.toggle('open');
}

// ── VALIDATE ──
function getFormData() {
  const apiKey = document.getElementById('apiKey').value.trim();
  const cvText = document.getElementById('cvText').value.trim();
  const jobDesc = document.getElementById('jobDesc').value.trim();
  const candidateName = document.getElementById('candidateName').value.trim() || 'Candidate';
  const companyName = document.getElementById('companyName').value.trim() || 'Target Company';
  const roleTitle = document.getElementById('roleTitle').value.trim() || 'Target Role';
  const topAchievement = document.getElementById('topAchievement').value.trim() || '';

  if (!apiKey) { showError('يرجى إدخال API Key'); return null; }
  if (!cvText && !jobDesc) { showError('يرجى إدخال السيرة الذاتية ووصف الوظيفة'); return null; }
  if (!jobDesc) { showError('يرجى إدخال وصف الوظيفة'); return null; }
  if (!cvText) { showError('يرجى إدخال نص السيرة الذاتية'); return null; }

  return { provider: selectedProvider, api_key: apiKey, cv_text: cvText, job_desc: jobDesc, candidate_name: candidateName, company_name: companyName, role_title: roleTitle, top_achievement: topAchievement };
}

// ── ANALYSIS ONLY ──
async function runAnalysis() {
  hideError();
  const fd = getFormData(); if (!fd) return;

  setLoader(true, 'جاري تحليل السيرة الذاتية...', 'مقارنة CV مع وصف الوظيفة');
  document.getElementById('results').style.display = 'none';

  try {
    const res = await fetch('/api/analyze', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify(fd)
    });
    const data = await res.json();
    if (!data.success) throw new Error(data.error);

    setLoader(false);
    generatedData.analysis = data.data;
    document.getElementById('results').style.display = 'block';
    document.getElementById('analysisCard').style.display = 'block';
    renderAnalysis(data.data);
  } catch(e) {
    setLoader(false);
    showError(e.message);
  }
}

// ── FULL GENERATION ──
async function runFullGeneration() {
  hideError();
  const fd = getFormData(); if (!fd) return;

  document.getElementById('results').style.display = 'none';
  setLoader(true, 'الذكاء الاصطناعي يعمل على ملفك...', 'جاري التحليل الشامل');
  [1,2,3,4,5,6,7].forEach(n => setDot(n, ''));

  const steps = [
    [1, 'تحليل السيرة الذاتية...'],
    [2, 'إعادة كتابة CV...'],
    [3, 'كتابة Cover Letter...'],
    [4, 'توليد رسائل LinkedIn...'],
    [5, 'بناء قصص STAR...'],
    [6, 'توليد أسئلة المقابلة...'],
    [7, 'بناء خطة ٣٠ يوم...']
  ];

  let stepIdx = 0;
  const stepTimer = setInterval(() => {
    if (stepIdx < steps.length) {
      const [n, txt] = steps[stepIdx];
      if (stepIdx > 0) setDot(stepIdx, 'done');
      setDot(n, 'active');
      document.getElementById('loaderStep').textContent = txt;
      stepIdx++;
    } else clearInterval(stepTimer);
  }, 3500);

  try {
    const res = await fetch('/api/generate_all', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify(fd)
    });
    const data = await res.json();
    clearInterval(stepTimer);

    if (!data.success) throw new Error(data.error);

    [1,2,3,4,5,6,7].forEach(n => setDot(n, 'done'));
    setLoader(false);

    generatedData = data.data;
    generatedData.candidate_name = fd.candidate_name;

    document.getElementById('results').style.display = 'block';
    document.getElementById('downloadBar').style.display = 'block';
    document.getElementById('outputTabs').style.display = 'block';
    document.getElementById('analysisCard').style.display = 'none';

    renderAllOutputs(data.data);
  } catch(e) {
    clearInterval(stepTimer);
    setLoader(false);
    showError(e.message);
  }
}

// ── RENDER ANALYSIS ──
function renderAnalysis(a) {
  const score = a.match_score || 0;
  const cls = score >= 75 ? 'score-excellent' : score >= 50 ? 'score-good' : 'score-fair';
  document.getElementById('scoreDisplay').innerHTML = `
    <div class="score-circle ${cls}">
      <div class="score-num">${score}</div>
      <div class="score-pct">% تطابق</div>
    </div>
    <p style="text-align:center;font-size:16px;font-weight:800;color:var(--blue);margin-bottom:16px">${a.match_label || ''}</p>
    <p style="text-align:center;font-size:13px;color:var(--gray);margin-bottom:20px;direction:ltr">${a.overall_verdict || ''}</p>
  `;

  let html = '';
  if (a.strengths?.length) {
    html += '<h3 style="font-size:14px;font-weight:700;color:var(--green);margin-bottom:10px">✅ نقاط القوة</h3><ul class="result-list">';
    a.strengths.forEach(s => { html += `<li class="list-strength"><div class="list-item-title">${s.title}</div>${s.detail}</li>`; });
    html += '</ul>';
  }
  if (a.weaknesses?.length) {
    html += '<h3 style="font-size:14px;font-weight:700;color:var(--red);margin:16px 0 10px">⚠️ نقاط الضعف والفجوات</h3><ul class="result-list">';
    a.weaknesses.forEach(w => { html += `<li class="list-weakness"><div class="list-item-title"><span class="badge badge-${(w.impact||'').toLowerCase()}">${w.impact}</span>${w.title}</div>${w.detail}</li>`; });
    html += '</ul>';
  }
  if (a.missing_keywords?.length) {
    html += '<h3 style="font-size:14px;font-weight:700;color:var(--orange);margin:16px 0 10px">🔑 كلمات مفتاحية ناقصة</h3><div class="keywords">';
    a.missing_keywords.forEach(k => { html += `<span class="keyword">${k}</span>`; });
    html += '</div>';
  }
  if (a.quick_wins?.length) {
    html += '<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:16px 0 10px">⚡ تحسينات سريعة</h3><ul class="result-list">';
    a.quick_wins.forEach(q => { html += `<li class="list-suggestion">${q}</li>`; });
    html += '</ul>';
  }
  document.getElementById('analysisDetails').innerHTML = html;
}

// ── RENDER ALL OUTPUTS ──
function renderAllOutputs(data) {
  // CV Tab
  if (data.cv) {
    const cv = data.cv;
    const h = cv.header || {};
    let cvHtml = `
      <div style="text-align:center;padding:16px;background:var(--blue);color:white;border-radius:12px;margin-bottom:16px">
        <div style="font-size:20px;font-weight:800">${h.name||''}</div>
        <div style="font-size:13px;margin-top:4px">${h.title||''}</div>
        <div style="font-size:11px;opacity:0.8;margin-top:4px">${h.phone||''} | ${h.email||''} | ${h.location||''}</div>
      </div>`;
    if (cv.summary) cvHtml += collapsibleSection('Profile Summary', cv.summary, '👤');
    if (cv.experience?.length) {
      let expHtml = '';
      cv.experience.forEach(exp => {
        expHtml += `<div style="margin-bottom:16px;padding:12px;background:var(--light);border-radius:8px;direction:ltr">
          <div style="font-weight:700;color:var(--blue)">${exp.title} — ${exp.company}</div>
          <div style="font-size:12px;color:var(--gray)">${exp.dates} | Reports to: ${exp.reports_to||''}</div>
          <ul style="margin-top:8px;padding-right:16px">
            ${(exp.bullets||[]).map(b=>`<li style="font-size:12px;margin-bottom:4px">${b}</li>`).join('')}
          </ul>
        </div>`;
      });
      cvHtml += collapsibleSection('الخبرة المهنية', expHtml, '💼');
    }
    if (cv.alignment_statement) cvHtml += collapsibleSection('Alignment Statement', cv.alignment_statement, '🎯');
    document.getElementById('cvOutput').innerHTML = cvHtml;
  }

  // Cover Letter Tab
  if (data.cover_letter) {
    const cl = data.cover_letter;
    let clHtml = `
      <div style="padding:16px;background:var(--blue);color:white;border-radius:12px;margin-bottom:16px">
        <div style="font-size:14px;font-weight:700">${cl.subject||''}</div>
      </div>
      <div style="direction:ltr;font-size:13px;line-height:1.8">
        <p style="margin-bottom:12px">Dear Hiring Manager,</p>
        <p style="margin-bottom:12px">${cl.opening||''}</p>`;
    if (cl.value_props?.length) {
      clHtml += '<div style="background:var(--blue-light);border-radius:8px;padding:12px;margin-bottom:12px"><strong style="color:var(--blue)">What I Bring:</strong><ul style="margin-top:8px;padding-right:20px">';
      cl.value_props.forEach(vp => { clHtml += `<li style="margin-bottom:4px">${vp.text}</li>`; });
      clHtml += '</ul></div>';
    }
    (cl.body_paragraphs||[]).forEach(p => { clHtml += `<p style="margin-bottom:12px">${p}</p>`; });
    clHtml += `<p>${cl.closing||''}</p>`;
    clHtml += `<p style="margin-top:16px;font-weight:700;color:var(--blue)">${(cl.signature||{}).name||''}</p>`;
    clHtml += '</div>';
    document.getElementById('clOutput').innerHTML = clHtml;
  }

  // LinkedIn Tab
  if (data.linkedin) {
    const li = data.linkedin;
    let liHtml = '';
    if (li.headline) liHtml += collapsibleSection('Recommended Headline', `<div style="direction:ltr;font-size:13px;background:var(--blue-light);padding:12px;border-radius:8px">${li.headline}</div>`, '📝');
    if (li.about_section) liHtml += collapsibleSection('About Section (Copy & Paste)', `<pre style="direction:ltr;font-size:12px;white-space:pre-wrap;line-height:1.6">${li.about_section}</pre>`, '👤');
    if (li.messages?.length) {
      li.messages.forEach(msg => {
        const content = (msg.subject ? `<strong>Subject:</strong> ${msg.subject}<br><br>` : '') + msg.body;
        liHtml += collapsibleSection(msg.type, `<div style="direction:ltr;font-size:13px;white-space:pre-wrap;line-height:1.6">${content}</div>`, '💬');
      });
    }
    document.getElementById('liOutput').innerHTML = liHtml;
  }

  // STAR Tab
  if (data.star?.stories) {
    let starHtml = '';
    data.star.stories.forEach((story, i) => {
      starHtml += `
        <div class="star-card">
          <div class="star-card-header">⭐ قصة ${i+1}: ${story.competency||''} — ${story.title||''}</div>
          <div class="star-row"><div class="star-label">SITUATION</div><div class="star-text">${story.situation||''}</div></div>
          <div class="star-row"><div class="star-label">TASK</div><div class="star-text">${story.task||''}</div></div>
          <div class="star-row"><div class="star-label">ACTION</div><div class="star-text">${story.action||''}</div></div>
          <div class="star-row star-result"><div class="star-label" style="color:var(--green)">RESULT</div><div class="star-text" style="font-weight:600">${story.result||''}</div></div>
          <div style="padding:10px 12px;font-size:12px;color:var(--blue);font-style:italic;direction:ltr">💡 ${story.key_message||''}</div>
        </div>`;
    });
    document.getElementById('starOutput').innerHTML = starHtml;
  }

  // Q&A Tab
  if (data.qa) {
    let qaHtml = '';
    (data.qa.questions||[]).forEach((qa, i) => {
      qaHtml += `
        <div style="margin-bottom:20px;border:1px solid var(--border);border-radius:12px;overflow:hidden">
          <div style="background:var(--blue);color:white;padding:12px 16px;font-size:13px;font-weight:700">❓ سؤال ${i+1}: ${qa.question||''}</div>
          <div style="padding:14px;font-size:13px;line-height:1.7;direction:ltr;background:var(--light)">${qa.answer||''}</div>
          <div style="padding:10px 14px;font-size:12px;color:var(--blue);background:var(--blue-light)">💡 ${qa.tip||''}</div>
        </div>`;
    });
    if (data.qa.questions_to_ask?.length) {
      qaHtml += '<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:20px 0 10px">❓ أسئلة تطرحها أنت على المحاور</h3>';
      data.qa.questions_to_ask.forEach(q => {
        qaHtml += `<div style="padding:12px;border-right:4px solid var(--blue);background:var(--blue-light);border-radius:0 8px 8px 0;margin-bottom:8px;font-size:13px;direction:ltr">${q}</div>`;
      });
    }
    document.getElementById('qaOutput').innerHTML = qaHtml;
  }

  // Plan Tab
  if (data.plan) {
    let planHtml = '';
    (data.plan.weeks||[]).forEach(week => {
      planHtml += `<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:16px 0 8px">${week.week||''}</h3>`;
      planHtml += `<p style="font-size:12px;color:var(--gray);margin-bottom:10px">${week.focus||''}</p>`;
      (week.tasks||[]).forEach(task => {
        const color = task.priority === 'HIGH' ? '#FEE2E2' : 'var(--light)';
        planHtml += `
          <div style="display:flex;gap:8px;padding:10px;border-radius:8px;background:${color};margin-bottom:6px;direction:ltr;align-items:flex-start">
            <span style="font-weight:700;font-size:11px;color:var(--blue);min-width:60px">${task.day||''}</span>
            <span style="font-size:12px;flex:1">${task.action||''}</span>
            ${task.platform ? `<span style="font-size:10px;background:white;border-radius:6px;padding:2px 6px;color:var(--gray)">${task.platform}</span>` : ''}
          </div>`;
      });
    });
    if (data.plan.pro_tips?.length) {
      planHtml += '<h3 style="font-size:14px;font-weight:700;color:var(--green);margin:20px 0 8px">💡 نصائح احترافية</h3>';
      data.plan.pro_tips.forEach(t => {
        planHtml += `<div style="padding:10px;border-right:4px solid var(--green);background:var(--green-light);border-radius:0 8px 8px 0;margin-bottom:6px;font-size:13px;direction:ltr">${t}</div>`;
      });
    }
    document.getElementById('planOutput').innerHTML = planHtml;
  }
}

// ── TAB SWITCH ──
function switchTab(tabId, el) {
  document.querySelectorAll('.tab-content').forEach(e => e.classList.remove('active'));
  document.querySelectorAll('.tab').forEach(e => e.classList.remove('active'));
  document.getElementById(tabId).classList.add('active');
  if (el) el.classList.add('active');
}

// ── DOWNLOAD ──
async function downloadDoc(type) {
  if (!generatedData || Object.keys(generatedData).length === 0) {
    alert('يرجى توليد الوثائق أولاً');
    return;
  }
  let body = {}, filename = '';
  if (type === 'cv') {
    body = { cv_data: generatedData.cv };
    filename = 'CV_Optimized.docx';
  } else if (type === 'cover_letter') {
    body = { cl_data: generatedData.cover_letter, candidate_info: { name: generatedData.candidate_name || '' } };
    filename = 'Cover_Letter.docx';
  } else if (type === 'strategy') {
    body = { linkedin_data: generatedData.linkedin, star_data: generatedData.star, qa_data: generatedData.qa, plan_data: generatedData.plan };
    filename = 'Job_Strategy_Pack.docx';
  }
  try {
    const res = await fetch('/api/download/' + type, {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify(body)
    });
    if (!res.ok) throw new Error('فشل التحميل');
    const blob = await res.blob();
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url; a.download = filename;
    document.body.appendChild(a); a.click();
    document.body.removeChild(a);
    URL.revokeObjectURL(url);
  } catch(e) {
    alert('خطأ في التحميل: ' + e.message);
  }
}
</script>
</body>
</html>"""


if __name__ == '__main__':
    print("""
╔══════════════════════════════════════════════════════════╗
║        CV ANALYZER PRO — مجاني ١٠٠٪                     ║
╠══════════════════════════════════════════════════════════╣
║  Local:   http://localhost:5000                         ║
║  Mobile:  http://YOUR_IP:5000  (same WiFi)             ║
╠══════════════════════════════════════════════════════════╣
║  Gemini Key:  aistudio.google.com/app/apikey           ║
║  Groq Key:    console.groq.com → API Keys              ║
╚══════════════════════════════════════════════════════════╝
""")
    app.run(host='0.0.0.0', port=5000, debug=False)
