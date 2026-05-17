"""
CV ANALYZER PRO — النسخة المُصلَحة
=====================================
الـ AI calls بتحصل من المتصفح مباشرة لـ Gemini/Groq
Render بيشيل DOCX generation بس — مش محتاج يكلم AI أصلاً

REQUIREMENTS:  flask  python-docx  pillow
BUILD COMMAND: pip install -r requirements.txt
START COMMAND: python app.py
"""

import io
from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)

# ─────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────

def hex_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading_line(doc, text, level=1, color_hex="1B4F8C"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    r, g, b = hex_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10)

# ─────────────────────────────────────────
# DOCX: CV
# ─────────────────────────────────────────

def generate_cv_docx(cv_data):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    h = cv_data.get("header", {})
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(h.get("name", "").upper())
    nr.bold = True; nr.font.size = Pt(22)
    nr.font.color.rgb = RGBColor(27, 79, 140)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(h.get("title", ""))
    tr.bold = True; tr.font.size = Pt(13)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f"{h.get('phone','')}  |  {h.get('email','')}  |  {h.get('location','')}")
    cr.font.size = Pt(10); cr.font.color.rgb = RGBColor(85, 85, 85)

    if h.get("education"):
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = ep.add_run(h["education"])
        er.font.size = Pt(10); er.font.color.rgb = RGBColor(85, 85, 85)

    if cv_data.get("summary"):
        add_heading_line(doc, "PROFESSIONAL SUMMARY")
        s = doc.add_paragraph(cv_data["summary"])
        for run in s.runs: run.font.size = Pt(10)

    if cv_data.get("competencies"):
        add_heading_line(doc, "CORE COMPETENCIES")
        table = doc.add_table(rows=len(cv_data["competencies"]), cols=2)
        table.style = 'Table Grid'
        for i, comp in enumerate(cv_data["competencies"]):
            row = table.rows[i]
            set_cell_bg(row.cells[0], "D6E4F0")
            row.cells[0].width = Inches(1.8)
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(comp.get("category", ""))
            r0.bold = True; r0.font.size = Pt(9)
            r0.font.color.rgb = RGBColor(27, 79, 140)
            row.cells[1].paragraphs[0].add_run(comp.get("skills", "")).font.size = Pt(9)

    if cv_data.get("career_highlights"):
        add_heading_line(doc, "CAREER HIGHLIGHTS")
        ch = cv_data["career_highlights"]
        ht = doc.add_paragraph()
        r = ht.add_run(ch.get("title", ""))
        r.bold = True; r.font.size = Pt(10)
        for b in ch.get("bullets", []): add_bullet(doc, b)

    if cv_data.get("experience"):
        add_heading_line(doc, "PROFESSIONAL EXPERIENCE")
        for exp in cv_data["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r1 = p.add_run(exp.get("title", "") + " — ")
            r1.bold = True; r1.font.size = Pt(11)
            r1.font.color.rgb = RGBColor(27, 79, 140)
            r2 = p.add_run(exp.get("company", ""))
            r2.bold = True; r2.font.size = Pt(11)
            r3 = p.add_run("  |  " + exp.get("location", ""))
            r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(85, 85, 85)
            dp2 = doc.add_paragraph()
            dr = dp2.add_run(exp.get("dates", "") + "  |  Reports to: " + exp.get("reports_to", ""))
            dr.italic = True; dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(85, 85, 85)
            dp2.paragraph_format.space_after = Pt(4)
            for b in exp.get("bullets", []): add_bullet(doc, b)

    if cv_data.get("education"):
        add_heading_line(doc, "EDUCATION & PROFESSIONAL DEVELOPMENT")
        for edu in cv_data["education"]:
            p = doc.add_paragraph()
            r = p.add_run(edu.get("degree", "") + " — " + edu.get("field", ""))
            r.bold = True; r.font.size = Pt(10)
            for b in edu.get("bullets", []): add_bullet(doc, b)

    if cv_data.get("alignment_statement"):
        add_heading_line(doc, "ALIGNMENT WITH TARGET ORGANIZATION")
        ap = doc.add_paragraph(cv_data["alignment_statement"])
        for run in ap.runs: run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ─────────────────────────────────────────
# DOCX: COVER LETTER
# ─────────────────────────────────────────

def generate_cover_letter_docx(cl_data, candidate_info):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85); section.right_margin = Inches(0.85)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(candidate_info.get("name", "").upper())
    nr.bold = True; nr.font.size = Pt(20)
    nr.font.color.rgb = RGBColor(27, 79, 140)

    sig = cl_data.get("signature", {})
    cp2 = doc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp2.add_run(f"{sig.get('phone','')}  |  {sig.get('email','')}  |  {sig.get('location','')}")
    cr.font.size = Pt(10); cr.font.color.rgb = RGBColor(85, 85, 85)

    div_p = doc.add_paragraph()
    pPr = div_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '1B4F8C')
    pBdr.append(bot); pPr.append(pBdr)

    doc.add_paragraph()
    sp = doc.add_paragraph()
    sr = sp.add_run(cl_data.get("subject", ""))
    sr.bold = True; sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(27, 79, 140)

    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")

    op = doc.add_paragraph(cl_data.get("opening", ""))
    op.paragraph_format.space_after = Pt(10)

    vps = cl_data.get("value_props", [])
    if vps:
        vp_label = doc.add_paragraph()
        vlr = vp_label.add_run("What I Bring to This Role:")
        vlr.bold = True; vlr.font.size = Pt(11)
        vlr.font.color.rgb = RGBColor(27, 79, 140)
        rows2 = [vps[i:i+2] for i in range(0, len(vps), 2)]
        table2 = doc.add_table(rows=len(rows2), cols=2)
        table2.style = 'Table Grid'
        for i, row_data in enumerate(rows2):
            row = table2.rows[i]
            for j, vp in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    if i % 2 == 0: set_cell_bg(cell, "D6E4F0")
                    r = cell.paragraphs[0].add_run(vp.get("check", "✓") + "  " + vp.get("text", ""))
                    r.font.size = Pt(9)

    doc.add_paragraph()
    for para_text in cl_data.get("body_paragraphs", []):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs: run.font.size = Pt(10)

    cp3 = doc.add_paragraph(cl_data.get("closing", ""))
    for run in cp3.runs: run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()

    np2 = doc.add_paragraph()
    nr2 = np2.add_run(sig.get("name", ""))
    nr2.bold = True; nr2.font.size = Pt(12)
    nr2.font.color.rgb = RGBColor(27, 79, 140)

    for field in ["title", "phone", "email", "location"]:
        fp2 = doc.add_paragraph(sig.get(field, ""))
        for run in fp2.runs:
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(85, 85, 85)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ─────────────────────────────────────────
# DOCX: STRATEGY PACK
# ─────────────────────────────────────────

def generate_strategy_docx(linkedin_data, star_data, qa_data, plan_data):
    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = tp.add_run("JOB APPLICATION STRATEGY PACK")
    tr2.bold = True; tr2.font.size = Pt(20)
    tr2.font.color.rgb = RGBColor(27, 79, 140)

    add_heading_line(doc, "SECTION 1 — LinkedIn Profile Optimization")
    if linkedin_data.get("headline"):
        tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "D6E4F0")
        tbl.rows[0].cells[0].paragraphs[0].add_run(linkedin_data["headline"]).font.size = Pt(10)
    if linkedin_data.get("about_section"):
        doc.add_paragraph()
        doc.add_paragraph().add_run("Recommended About Section:").bold = True
        tbl2 = doc.add_table(rows=1, cols=1); tbl2.style = 'Table Grid'
        set_cell_bg(tbl2.rows[0].cells[0], "F5F5F5")
        tbl2.rows[0].cells[0].paragraphs[0].add_run(linkedin_data["about_section"]).font.size = Pt(9)
    for tip in linkedin_data.get("profile_tips", []): add_bullet(doc, tip)

    add_heading_line(doc, "SECTION 2 — LinkedIn Outreach Messages")
    for msg in linkedin_data.get("messages", []):
        doc.add_paragraph()
        doc.add_paragraph().add_run(f"Message: {msg.get('type','')}").bold = True
        tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "EBF5FB")
        body_text = (f"Subject: {msg['subject']}\n\n" if msg.get('subject') else "") + msg.get("body", "")
        tbl.rows[0].cells[0].paragraphs[0].add_run(body_text).font.size = Pt(9)

    add_heading_line(doc, "SECTION 3 — STAR Interview Stories")
    for story in star_data.get("stories", []):
        doc.add_paragraph()
        sh = doc.add_paragraph()
        sr2 = sh.add_run(f"Competency: {story.get('competency','')} — {story.get('title','')}")
        sr2.bold = True; sr2.font.color.rgb = RGBColor(27, 79, 140)
        tbl = doc.add_table(rows=4, cols=2); tbl.style = 'Table Grid'
        for i, (lbl, key, col) in enumerate(zip(
            ["SITUATION","TASK","ACTION","RESULT"],
            ["situation","task","action","result"],
            ["D6E4F0","F5F5F5","D6E4F0","D6F0E0"]
        )):
            row = tbl.rows[i]
            set_cell_bg(row.cells[0], col)
            row.cells[0].paragraphs[0].add_run(lbl).bold = True
            row.cells[1].paragraphs[0].add_run(story.get(key, "")).font.size = Pt(9)
        km = doc.add_paragraph()
        kmr = km.add_run("Key Message: " + story.get("key_message", ""))
        kmr.italic = True; kmr.font.size = Pt(9)
        kmr.font.color.rgb = RGBColor(27, 79, 140)

    add_heading_line(doc, "SECTION 4 — Interview Questions & Answers")
    for qa in qa_data.get("questions", []):
        doc.add_paragraph()
        qp = doc.add_paragraph()
        qr2 = qp.add_run("Q: " + qa.get("question", ""))
        qr2.bold = True; qr2.font.color.rgb = RGBColor(27, 79, 140)
        tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], "F0F8FF")
        tbl.rows[0].cells[0].paragraphs[0].add_run(qa.get("answer", "")).font.size = Pt(9)
        tip_p = doc.add_paragraph()
        tip_r = tip_p.add_run("💡 " + qa.get("tip", ""))
        tip_r.italic = True; tip_r.font.size = Pt(9)
        tip_r.font.color.rgb = RGBColor(85, 85, 85)
    if qa_data.get("questions_to_ask"):
        add_heading_line(doc, "Questions to Ask the Interviewer", level=2)
        for q in qa_data["questions_to_ask"]: add_bullet(doc, q)

    add_heading_line(doc, "SECTION 5 — 30-Day Action Plan")
    for week in plan_data.get("weeks", []):
        doc.add_paragraph()
        wh = doc.add_paragraph()
        wr2 = wh.add_run(week.get("week", ""))
        wr2.bold = True; wr2.font.color.rgb = RGBColor(27, 79, 140)
        fp2 = doc.add_paragraph(week.get("focus", ""))
        if fp2.runs: fp2.runs[0].italic = True
        for task in week.get("tasks", []):
            p = doc.add_paragraph(style='List Bullet')
            r1 = p.add_run(f"[{task.get('day','')}] ")
            r1.bold = True; r1.font.size = Pt(9)
            r2 = p.add_run(task.get("action", ""))
            r2.font.size = Pt(9)
            if task.get("platform"):
                r3 = p.add_run(f" [{task['platform']}]")
                r3.font.size = Pt(8); r3.font.color.rgb = RGBColor(85, 85, 85)
    if plan_data.get("pro_tips"):
        add_heading_line(doc, "Pro Tips", level=2)
        for tip in plan_data["pro_tips"]: add_bullet(doc, "💡 " + tip)
    if plan_data.get("red_flags_to_avoid"):
        add_heading_line(doc, "Common Mistakes to Avoid", level=2)
        for flag in plan_data["red_flags_to_avoid"]: add_bullet(doc, "⚠️ " + flag)

    doc.add_paragraph()
    fp3 = doc.add_paragraph()
    fp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp3.add_run("You are not just a candidate. You are the solution they are looking for.")
    fr.bold = True; fr.font.size = Pt(13)
    fr.font.color.rgb = RGBColor(27, 79, 140)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ─────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/api/download/<doc_type>', methods=['POST'])
def api_download(doc_type):
    try:
        data = request.json
        if doc_type == 'cv':
            buf = generate_cv_docx(data['cv_data'])
            filename = "CV_Optimized.docx"
        elif doc_type == 'cover_letter':
            buf = generate_cover_letter_docx(data['cl_data'], data.get('candidate_info', {}))
            filename = "Cover_Letter.docx"
        elif doc_type == 'strategy':
            buf = generate_strategy_docx(
                data['linkedin_data'], data['star_data'],
                data['qa_data'], data['plan_data'])
            filename = "Job_Strategy_Pack.docx"
        else:
            return jsonify({"error": "Unknown doc type"}), 400
        return send_file(buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─────────────────────────────────────────
# HTML TEMPLATE  (AI calls من المتصفح مباشرة)
# ─────────────────────────────────────────

HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="ar" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1">
<title>CV Analyzer Pro</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;700;800&family=Space+Mono&display=swap');
:root{--blue:#1B4F8C;--blue-light:#D6E4F0;--blue-mid:#2E6DB4;--green:#1A6B3C;--green-light:#D6F0E0;--red:#B91C1C;--orange:#C0550A;--dark:#0F1923;--gray:#6B7280;--light:#F8FAFC;--border:#E2E8F0;--shadow:0 4px 24px rgba(27,79,140,0.12);--radius:16px}
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Tajawal',sans-serif;background:linear-gradient(135deg,#0F1923 0%,#1B4F8C 50%,#0F1923 100%);min-height:100vh;color:var(--dark)}
.hdr{background:rgba(255,255,255,0.05);backdrop-filter:blur(20px);border-bottom:1px solid rgba(255,255,255,0.1);padding:14px 20px;display:flex;align-items:center;gap:12px;position:sticky;top:0;z-index:100}
.logo{width:42px;height:42px;background:var(--blue);border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:20px;flex-shrink:0}
.ttl{color:white;flex:1}.ttl h1{font-size:17px;font-weight:800}.ttl p{font-size:11px;opacity:.6}
.badge{background:linear-gradient(135deg,#22c55e,#16a34a);color:white;font-size:11px;font-weight:800;padding:4px 10px;border-radius:20px}
.main{max-width:900px;margin:0 auto;padding:20px 16px 40px}
.card{background:white;border-radius:var(--radius);padding:24px;margin-bottom:16px;box-shadow:var(--shadow)}
.ch{display:flex;align-items:center;gap:10px;margin-bottom:18px}
.ci{width:38px;height:38px;background:var(--blue-light);border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:18px;flex-shrink:0}
.ct{font-size:16px;font-weight:700;color:var(--blue)}.cs{font-size:12px;color:var(--gray)}
label{font-size:13px;font-weight:600;color:var(--dark);margin-bottom:6px;display:block}
input[type=text],input[type=password],textarea{width:100%;padding:12px 14px;border:2px solid var(--border);border-radius:10px;font-family:'Tajawal',sans-serif;font-size:14px;color:var(--dark);transition:border-color .2s;margin-bottom:14px;background:var(--light);direction:ltr}
input:focus,textarea:focus{outline:none;border-color:var(--blue);background:white}
textarea{min-height:120px;resize:vertical}
.row{display:grid;grid-template-columns:1fr 1fr;gap:12px}
@media(max-width:600px){.row{grid-template-columns:1fr}}
.pvbox{background:linear-gradient(135deg,#0F1923,#1B3A6B);border-radius:16px;padding:20px;margin-bottom:16px;color:white}
.pvbox h3{font-size:14px;font-weight:800;margin-bottom:4px}
.pvbox p{font-size:12px;opacity:.7;margin-bottom:16px}
.ptabs{display:flex;gap:8px;margin-bottom:16px}
.ptab{flex:1;padding:12px 8px;border-radius:12px;border:2px solid rgba(255,255,255,.15);background:rgba(255,255,255,.05);color:rgba(255,255,255,.6);cursor:pointer;text-align:center;font-family:'Tajawal',sans-serif;font-size:13px;font-weight:700;transition:all .2s}
.ptab.ag{border-color:#4285F4;background:rgba(66,133,244,.15);color:#93C5FD}
.ptab.aq{border-color:#FF6B35;background:rgba(255,107,53,.15);color:#FCA5A5}
.ptab-logo{font-size:20px;display:block;margin-bottom:4px}
.ptab-sub{font-size:10px;opacity:.8;font-weight:400}
.akbox input{background:rgba(255,255,255,.08)!important;border-color:rgba(255,255,255,.15)!important;color:white!important;margin-bottom:6px}
.akbox input::placeholder{color:rgba(255,255,255,.35)!important}
.akbox label{color:rgba(255,255,255,.8)}
.alink{color:#60A5FA;font-size:12px;text-decoration:underline}
.ahint{font-size:12px;color:rgba(255,255,255,.55)}
.btn{display:inline-flex;align-items:center;gap:8px;padding:14px 24px;border-radius:12px;font-size:14px;font-weight:700;font-family:'Tajawal',sans-serif;cursor:pointer;border:none;transition:all .2s}
.btn-p{background:var(--blue);color:white;width:100%;justify-content:center;font-size:15px;padding:16px}
.btn-p:hover{background:var(--blue-mid);transform:translateY(-1px)}
.btn-p:disabled{opacity:.5;cursor:not-allowed;transform:none}
.btn-o{background:white;color:var(--blue);border:2px solid var(--blue);width:100%;justify-content:center;margin-bottom:10px}
.btn-o:hover{background:var(--blue-light)}
.tabs{display:flex;gap:6px;margin-bottom:20px;flex-wrap:wrap}
.tab{padding:8px 16px;border-radius:20px;font-size:13px;font-weight:600;cursor:pointer;border:2px solid var(--border);background:white;color:var(--gray);transition:all .2s;font-family:'Tajawal',sans-serif}
.tab.active{background:var(--blue);color:white;border-color:var(--blue)}
.tc{display:none}.tc.active{display:block}
.sc{width:120px;height:120px;border-radius:50%;display:flex;flex-direction:column;align-items:center;justify-content:center;margin:0 auto 16px}
.sn{font-size:42px;font-weight:800;line-height:1;font-family:'Space Mono',monospace}
.sp{font-size:14px;font-weight:600}
.se{background:var(--green-light);color:var(--green);border:4px solid var(--green)}
.sg{background:#FEF9C3;color:#854D0E;border:4px solid #CA8A04}
.sf{background:#FEE2E2;color:var(--red);border:4px solid var(--red)}
.rl{list-style:none}
.rl li{padding:12px 14px;border-radius:10px;margin-bottom:8px;font-size:13px;line-height:1.5;direction:ltr}
.ls{background:var(--green-light);border-right:4px solid var(--green)}
.lw{background:#FEE2E2;border-right:4px solid var(--red)}
.li{background:var(--blue-light);border-right:4px solid var(--blue)}
.lit{font-weight:700;margin-bottom:3px}
.bdg{display:inline-block;padding:2px 8px;border-radius:20px;font-size:10px;font-weight:700;margin-right:6px}
.bh{background:#FEE2E2;color:var(--red)}.bm{background:#FEF9C3;color:#854D0E}.bl{background:var(--green-light);color:var(--green)}
.kws{display:flex;flex-wrap:wrap;gap:8px;margin-top:10px}
.kw{padding:4px 12px;border-radius:20px;font-size:12px;font-weight:600;background:#FEE2E2;color:var(--red);border:1px solid #FECACA}
.ldr{display:none;text-align:center;padding:40px 20px}
.ldr.show{display:block}
.spin{width:50px;height:50px;border:4px solid var(--blue-light);border-top-color:var(--blue);border-radius:50%;animation:spin .8s linear infinite;margin:0 auto 16px}
@keyframes spin{to{transform:rotate(360deg)}}
.lt{font-size:14px;color:var(--gray)}.ls2{font-size:12px;color:var(--blue);margin-top:6px;font-weight:600}
.psteps{display:flex;gap:8px;margin-top:20px;overflow-x:auto;padding-bottom:4px}
.pstep{display:flex;flex-direction:column;align-items:center;min-width:65px}
.pdot{width:32px;height:32px;border-radius:50%;background:var(--border);color:var(--gray);display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:700;margin-bottom:4px;transition:all .3s}
.pdot.done{background:var(--green);color:white}.pdot.active{background:var(--blue);color:white;animation:pulse 1s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.7}}
.plbl{font-size:9px;color:var(--gray);text-align:center}
.os{border:1px solid var(--border);border-radius:12px;overflow:hidden;margin-bottom:16px}
.osh{background:var(--blue);color:white;padding:12px 16px;font-size:14px;font-weight:700;display:flex;align-items:center;gap:8px;cursor:pointer}
.osb{padding:16px;font-size:13px;line-height:1.7;direction:ltr;white-space:pre-wrap;background:var(--light);display:none}
.osb.open{display:block}
.dlbar{background:linear-gradient(135deg,var(--blue),var(--blue-mid));border-radius:16px;padding:20px;color:white;margin-bottom:16px}
.dlbar h3{font-size:16px;font-weight:800;margin-bottom:14px}
.dlbtns{display:flex;gap:10px;flex-wrap:wrap}
.dlbtn{background:white;color:var(--blue);border:none;border-radius:10px;padding:10px 18px;font-size:13px;font-weight:700;cursor:pointer;display:flex;align-items:center;gap:6px;font-family:'Tajawal',sans-serif;transition:all .2s}
.dlbtn:hover{transform:translateY(-2px);box-shadow:0 4px 12px rgba(0,0,0,.15)}
.alert{padding:12px 16px;border-radius:10px;font-size:13px;margin-bottom:14px;display:none}
.alert.show{display:block}
.ae{background:#FEE2E2;color:var(--red);border:1px solid #FECACA}
.stc{border:1px solid var(--border);border-radius:12px;overflow:hidden;margin-bottom:16px}
.stch{background:var(--blue);color:white;padding:12px 16px;font-size:14px;font-weight:700}
.str{display:grid;grid-template-columns:95px 1fr;border-bottom:1px solid var(--border)}
.str:last-child{border-bottom:none}
.stl{background:var(--blue-light);color:var(--blue);font-weight:700;font-size:12px;padding:12px;display:flex;align-items:center}
.stt{padding:12px;font-size:13px;line-height:1.6;direction:ltr}
.str-r{background:var(--green-light)}
@media(max-width:480px){.card{padding:16px}.sc{width:90px;height:90px}.sn{font-size:32px}}
</style>
</head>
<body>

<div class="hdr">
  <div class="logo">🎯</div>
  <div class="ttl"><h1>CV Analyzer Pro</h1><p>محلل السيرة الذاتية الذكي</p></div>
  <span class="badge">FREE ∞</span>
</div>

<div class="main">

  <div class="pvbox">
    <h3>🤖 اختر مزود الذكاء الاصطناعي</h3>
    <p>كلاهم مجاني تماماً — لا بطاقة ائتمان — للأبد</p>
    <div class="ptabs">
      <div class="ptab ag" id="tabG" onclick="selP('gemini')">
        <span class="ptab-logo">🔵</span>Google Gemini
        <div class="ptab-sub">مجاني — 15 req/min</div>
      </div>
      <div class="ptab" id="tabQ" onclick="selP('groq')">
        <span class="ptab-logo">🟠</span>Groq
        <div class="ptab-sub">مجاني — سريع جداً</div>
      </div>
    </div>
    <div class="akbox">
      <label id="kLbl">🔑 Gemini API Key</label>
      <input type="password" id="apiKey" placeholder="AIza..." autocomplete="off">
      <p class="ahint" id="kHint">احصل على مفتاح مجاني: <a class="alink" href="https://aistudio.google.com/app/apikey" target="_blank">aistudio.google.com</a> — سجّل بحساب Google</p>
    </div>
  </div>

  <div class="card">
    <div class="ch"><div class="ci">📋</div><div><div class="ct">أدخل البيانات</div><div class="cs">السيرة الذاتية + وصف الوظيفة</div></div></div>
    <div class="row">
      <div><label>اسمك الكامل</label><input type="text" id="cName" placeholder="Haitham El Meslemani"></div>
      <div><label>الشركة المستهدفة</label><input type="text" id="cCompany" placeholder="M42 / HealthPlus"></div>
    </div>
    <div class="row">
      <div><label>المسمى الوظيفي</label><input type="text" id="cRole" placeholder="Senior Operations Manager"></div>
      <div><label>أبرز إنجاز</label><input type="text" id="cAch" placeholder="Scaled clinic from 1.2M to 55M AED"></div>
    </div>
    <label>نص السيرة الذاتية</label>
    <textarea id="cvText" placeholder="الصق نص سيرتك الذاتية هنا..."></textarea>
    <label>وصف الوظيفة</label>
    <textarea id="jobDesc" placeholder="الصق وصف الوظيفة هنا..."></textarea>
    <div id="errBox" class="alert ae">⚠️ <span id="errMsg"></span></div>
    <button class="btn btn-o" onclick="doAnalysis()">🔍 تحليل التطابق فقط (سريع)</button>
    <button class="btn btn-p" onclick="doFull()">⚡ توليد كل الوثائق دفعة واحدة</button>
  </div>

  <div id="ldr" class="card ldr">
    <div class="spin"></div>
    <div class="lt" id="lTxt">جاري العمل...</div>
    <div class="ls2" id="lStep">يرجى الانتظار</div>
    <div class="psteps">
      <div class="pstep"><div class="pdot" id="d1">1</div><div class="plbl">تحليل</div></div>
      <div class="pstep"><div class="pdot" id="d2">2</div><div class="plbl">إعادة كتابة</div></div>
      <div class="pstep"><div class="pdot" id="d3">3</div><div class="plbl">Cover Letter</div></div>
      <div class="pstep"><div class="pdot" id="d4">4</div><div class="plbl">LinkedIn</div></div>
      <div class="pstep"><div class="pdot" id="d5">5</div><div class="plbl">STAR</div></div>
      <div class="pstep"><div class="pdot" id="d6">6</div><div class="plbl">مقابلة</div></div>
      <div class="pstep"><div class="pdot" id="d7">7</div><div class="plbl">خطة ٣٠ يوم</div></div>
    </div>
  </div>

  <div id="res" style="display:none">

    <div id="anCard" class="card" style="display:none">
      <div class="ch"><div class="ci">📊</div><div><div class="ct">نتيجة التحليل</div><div class="cs">CV vs Job Description</div></div></div>
      <div id="scoreDiv"></div>
      <div id="detailDiv"></div>
    </div>

    <div id="dlBar" class="dlbar" style="display:none">
      <h3>⬇️ تحميل الوثائق</h3>
      <div class="dlbtns">
        <button class="dlbtn" onclick="dlDoc('cv')">📄 CV المُحسَّنة</button>
        <button class="dlbtn" onclick="dlDoc('cover_letter')">✉️ Cover Letter</button>
        <button class="dlbtn" onclick="dlDoc('strategy')">🎯 حزمة الاستراتيجية</button>
      </div>
    </div>

    <div id="outTabs" class="card" style="display:none">
      <div class="tabs">
        <div class="tab active" onclick="swTab('tCV',this)">📄 السيرة</div>
        <div class="tab" onclick="swTab('tCL',this)">✉️ Cover Letter</div>
        <div class="tab" onclick="swTab('tLI',this)">💼 LinkedIn</div>
        <div class="tab" onclick="swTab('tST',this)">⭐ STAR</div>
        <div class="tab" onclick="swTab('tQA',this)">🎤 مقابلة</div>
        <div class="tab" onclick="swTab('tPL',this)">📅 خطة ٣٠ يوم</div>
      </div>
      <div id="tCV" class="tc active"><div id="cvOut"></div></div>
      <div id="tCL" class="tc"><div id="clOut"></div></div>
      <div id="tLI" class="tc"><div id="liOut"></div></div>
      <div id="tST" class="tc"><div id="stOut"></div></div>
      <div id="tQA" class="tc"><div id="qaOut"></div></div>
      <div id="tPL" class="tc"><div id="plOut"></div></div>
    </div>
  </div>
</div>

<script>
var prov='gemini', gd={};

function selP(p){
  prov=p;
  document.getElementById('tabG').className='ptab'+(p==='gemini'?' ag':'');
  document.getElementById('tabQ').className='ptab'+(p==='groq'?' aq':'');
  var inp=document.getElementById('apiKey');
  if(p==='gemini'){
    document.getElementById('kLbl').textContent='🔑 Gemini API Key';
    inp.placeholder='AIza...';
    document.getElementById('kHint').innerHTML='احصل على مفتاح مجاني: <a class="alink" href="https://aistudio.google.com/app/apikey" target="_blank">aistudio.google.com</a> — سجّل بحساب Google';
  } else {
    document.getElementById('kLbl').textContent='🔑 Groq API Key';
    inp.placeholder='gsk_...';
    document.getElementById('kHint').innerHTML='احصل على مفتاح مجاني: <a class="alink" href="https://console.groq.com" target="_blank">console.groq.com</a> ← API Keys ← Create Key';
  }
  inp.value='';
}

function showErr(m){document.getElementById('errMsg').textContent=m;document.getElementById('errBox').classList.add('show');}
function hideErr(){document.getElementById('errBox').classList.remove('show');}
function setLdr(show,txt,step){
  document.getElementById('ldr').classList.toggle('show',show);
  if(txt)document.getElementById('lTxt').textContent=txt;
  if(step)document.getElementById('lStep').textContent=step;
}
function setDot(n,s){
  var el=document.getElementById('d'+n);if(!el)return;
  el.className='pdot'+(s==='done'?' done':s==='active'?' active':'');
  if(s==='done')el.textContent='✓';
}
function togSec(id){document.getElementById(id).classList.toggle('open');}
function colSec(title,content,icon){
  var id='s'+Math.random().toString(36).slice(2,8);
  return '<div class="os"><div class="osh" onclick="togSec(\''+id+'\')">'+(icon?icon+' ':'')+title+' <span style="margin-right:auto;font-size:11px;opacity:.6">▼</span></div><div id="'+id+'" class="osb">'+content+'</div></div>';
}
function swTab(id,el){
  document.querySelectorAll('.tc').forEach(function(e){e.classList.remove('active');});
  document.querySelectorAll('.tab').forEach(function(e){e.classList.remove('active');});
  document.getElementById(id).classList.add('active');
  if(el)el.classList.add('active');
}

async function callAI(prompt,jsonMode){
  var key=document.getElementById('apiKey').value.trim();
  if(!key)throw new Error('يرجى إدخال API Key');
  var fp=jsonMode?prompt+'\n\nCRITICAL: Return ONLY valid JSON. No markdown, no backticks, no explanation.':prompt;

  if(prov==='gemini'){
    var models=['gemini-2.0-flash','gemini-1.5-flash-latest','gemini-1.5-flash-8b-latest'];
    var lastErr;
    for(var i=0;i<models.length;i++){
      try{
        var url='https://generativelanguage.googleapis.com/v1beta/models/'+models[i]+':generateContent?key='+key;
        var res=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json'},
          body:JSON.stringify({contents:[{parts:[{text:fp}]}],generationConfig:{maxOutputTokens:8192,temperature:0.7}})});
        var d=await res.json();
        if(!res.ok){lastErr=(d.error&&d.error.message)||JSON.stringify(d.error);continue;}
        var txt=d.candidates&&d.candidates[0]&&d.candidates[0].content&&d.candidates[0].content.parts&&d.candidates[0].content.parts[0]&&d.candidates[0].content.parts[0].text||'';
        if(jsonMode){txt=txt.replace(/^```(?:json)?\s*/,'').replace(/\s*```$/,'').trim();return JSON.parse(txt);}
        return txt;
      }catch(e){lastErr=e.message;}
    }
    throw new Error('Gemini Error: '+lastErr);
  } else {
    var models2=['llama-3.3-70b-versatile','llama-3.1-8b-instant','mixtral-8x7b-32768'];
    var lastErr2;
    for(var j=0;j<models2.length;j++){
      try{
        var res2=await fetch('https://api.groq.com/openai/v1/chat/completions',{method:'POST',
          headers:{'Content-Type':'application/json','Authorization':'Bearer '+key},
          body:JSON.stringify({model:models2[j],max_tokens:8192,temperature:0.7,
            messages:[{role:'system',content:'You are an expert career coach and HR consultant. UAE healthcare sector specialist.'},
                      {role:'user',content:fp}]})});
        var d2=await res2.json();
        if(!res2.ok){lastErr2=(d2.error&&d2.error.message)||JSON.stringify(d2.error);continue;}
        var txt2=d2.choices&&d2.choices[0]&&d2.choices[0].message&&d2.choices[0].message.content||'';
        if(jsonMode){txt2=txt2.replace(/^```(?:json)?\s*/,'').replace(/\s*```$/,'').trim();return JSON.parse(txt2);}
        return txt2;
      }catch(e){lastErr2=e.message;}
    }
    throw new Error('Groq Error: '+lastErr2);
  }
}

function getF(){
  var k=document.getElementById('apiKey').value.trim();
  var cv=document.getElementById('cvText').value.trim();
  var jd=document.getElementById('jobDesc').value.trim();
  if(!k){showErr('يرجى إدخال API Key');return null;}
  if(!cv){showErr('يرجى إدخال نص السيرة الذاتية');return null;}
  if(!jd){showErr('يرجى إدخال وصف الوظيفة');return null;}
  return{cv:cv,jd:jd,name:document.getElementById('cName').value.trim()||'Candidate',
    company:document.getElementById('cCompany').value.trim()||'Target Company',
    role:document.getElementById('cRole').value.trim()||'Target Role',
    ach:document.getElementById('cAch').value.trim()||''};
}

async function doAnalysis(){
  hideErr();var f=getF();if(!f)return;
  setLdr(true,'جاري تحليل السيرة...','مقارنة CV مع وصف الوظيفة');
  document.getElementById('res').style.display='none';
  try{
    setDot(1,'active');
    var r=await callAI('You are a senior HR consultant for UAE healthcare.\nAnalyze this CV against the JD.\n\nJD:\n'+f.jd+'\n\nCV:\n'+f.cv+'\n\nReturn JSON:\n{"match_score":<0-100>,"match_label":"<Poor|Fair|Good|Strong|Excellent>","strengths":[{"title":"...","detail":"..."}],"weaknesses":[{"title":"...","detail":"...","impact":"High|Medium|Low"}],"missing_keywords":["..."],"quick_wins":["..."],"overall_verdict":"..."}',true);
    setDot(1,'done');setLdr(false);
    gd.analysis=r;
    document.getElementById('res').style.display='block';
    document.getElementById('anCard').style.display='block';
    renderAn(r);
  }catch(e){setLdr(false);showErr(e.message);}
}

async function doFull(){
  hideErr();var f=getF();if(!f)return;
  document.getElementById('res').style.display='none';
  setLdr(true,'الذكاء الاصطناعي يعمل على ملفك...','');
  [1,2,3,4,5,6,7].forEach(function(n){setDot(n,'');});

  var steps=[
    [1,'تحليل التطابق...','You are a senior HR consultant for UAE healthcare.\nAnalyze this CV against the JD.\n\nJD:\n'+f.jd+'\n\nCV:\n'+f.cv+'\n\nReturn JSON:\n{"match_score":<0-100>,"match_label":"<Poor|Fair|Good|Strong|Excellent>","strengths":[{"title":"...","detail":"..."}],"weaknesses":[{"title":"...","detail":"...","impact":"High|Medium|Low"}],"missing_keywords":["..."],"quick_wins":["..."],"overall_verdict":"..."}'],
    [2,'إعادة كتابة السيرة...','You are an expert CV writer for UAE healthcare executives.\nRewrite this CV to 95%+ match. Candidate: '+f.name+'\n\nJD: '+f.jd+'\nCV: '+f.cv+'\n\nRules: Keep ALL real facts. Match JD keywords. Add measurable results.\n\nReturn JSON:\n{"header":{"name":"...","title":"...","phone":"...","email":"...","location":"...","education":"..."},"summary":"...","competencies":[{"category":"...","skills":"skill1 | skill2"}],"career_highlights":{"title":"...","bullets":["..."]},"experience":[{"title":"...","company":"...","location":"...","dates":"...","reports_to":"...","bullets":["..."]}],"education":[{"degree":"...","field":"...","bullets":["..."]}],"alignment_statement":"..."}'],
    [3,'كتابة Cover Letter...','Write a powerful cover letter.\nCandidate: '+f.name+' | Role: '+f.role+' | Company: '+f.company+'\nJD: '+f.jd+'\nCV: '+f.cv.slice(0,2000)+'\n\nReturn JSON:\n{"subject":"...","opening":"...","value_props":[{"check":"✓","text":"..."}],"body_paragraphs":["...","...","..."],"closing":"...","signature":{"name":"...","title":"...","phone":"...","email":"...","location":"..."}}'],
    [4,'توليد رسائل LinkedIn...','Generate LinkedIn content for '+f.name+' applying to '+f.role+' at '+f.company+'.\nAchievement: '+f.ach+'\n\nReturn JSON:\n{"headline":"...","about_section":"...","messages":[{"type":"Hiring Manager","subject":"...","body":"..."},{"type":"Connection Request (300 chars)","subject":null,"body":"..."},{"type":"Internal Employee Intel","subject":null,"body":"..."}],"profile_tips":["...","...","..."]}'],
    [5,'بناء قصص STAR...','Generate 5 STAR interview stories.\nJD: '+f.jd+'\nCV: '+f.cv.slice(0,3000)+'\n\nReturn JSON:\n{"stories":[{"competency":"...","title":"...","situation":"...","task":"...","action":"...","result":"...","key_message":"..."}]}'],
    [6,'توليد أسئلة المقابلة...','Generate 6 interview Q&As for role at '+f.company+'.\nJD: '+f.jd+'\nCV: '+f.cv.slice(0,2000)+'\n\nReturn JSON:\n{"questions":[{"question":"...","answer":"...","tip":"..."}],"questions_to_ask":["...","...","..."]}'],
    [7,'بناء خطة ٣٠ يوم...','Create a 30-day action plan for '+f.name+' to secure '+f.role+' at '+f.company+'.\n\nReturn JSON:\n{"weeks":[{"week":"Week 1 — Apply & Activate","focus":"...","tasks":[{"day":"Day 1-2","action":"...","platform":"LinkedIn|Email|Phone|Research","priority":"HIGH|MEDIUM"}]}],"pro_tips":["..."],"red_flags_to_avoid":["..."]}'],
  ];
  var keys=['analysis','cv','cover_letter','linkedin','star','qa','plan'];

  try{
    for(var i=0;i<steps.length;i++){
      if(i>0)setDot(i,'done');
      setDot(steps[i][0],'active');
      document.getElementById('lStep').textContent=steps[i][1];
      gd[keys[i]]=await callAI(steps[i][2],true);
    }
    setDot(7,'done');gd.cName=f.name;setLdr(false);
    document.getElementById('res').style.display='block';
    document.getElementById('dlBar').style.display='block';
    document.getElementById('outTabs').style.display='block';
    document.getElementById('anCard').style.display='block';
    renderAn(gd.analysis);renderAll(gd);
  }catch(e){setLdr(false);showErr(e.message);}
}

function renderAn(a){
  var sc=a.match_score||0;
  var cls=sc>=75?'sc se':sc>=50?'sc sg':'sc sf';
  document.getElementById('scoreDiv').innerHTML='<div class="'+cls+'"><div class="sn">'+sc+'</div><div class="sp">% تطابق</div></div><p style="text-align:center;font-size:16px;font-weight:800;color:var(--blue);margin-bottom:12px">'+( a.match_label||'')+'</p><p style="text-align:center;font-size:13px;color:var(--gray);margin-bottom:20px;direction:ltr">'+(a.overall_verdict||'')+'</p>';
  var h='';
  if(a.strengths&&a.strengths.length){h+='<h3 style="font-size:14px;font-weight:700;color:var(--green);margin-bottom:10px">✅ نقاط القوة</h3><ul class="rl">';a.strengths.forEach(function(s){h+='<li class="ls"><div class="lit">'+s.title+'</div>'+s.detail+'</li>';});h+='</ul>';}
  if(a.weaknesses&&a.weaknesses.length){h+='<h3 style="font-size:14px;font-weight:700;color:var(--red);margin:16px 0 10px">⚠️ الفجوات</h3><ul class="rl">';a.weaknesses.forEach(function(w){h+='<li class="lw"><div class="lit"><span class="bdg b'+((w.impact||'').charAt(0).toLowerCase())+'">'+w.impact+'</span>'+w.title+'</div>'+w.detail+'</li>';});h+='</ul>';}
  if(a.missing_keywords&&a.missing_keywords.length){h+='<h3 style="font-size:14px;font-weight:700;color:var(--orange);margin:16px 0 10px">🔑 كلمات ناقصة</h3><div class="kws">';a.missing_keywords.forEach(function(k){h+='<span class="kw">'+k+'</span>';});h+='</div>';}
  if(a.quick_wins&&a.quick_wins.length){h+='<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:16px 0 10px">⚡ تحسينات سريعة</h3><ul class="rl">';a.quick_wins.forEach(function(q){h+='<li class="li">'+q+'</li>';});h+='</ul>';}
  document.getElementById('detailDiv').innerHTML=h;
}

function renderAll(d){
  if(d.cv){
    var h2=d.cv.header||{};
    var html='<div style="text-align:center;padding:16px;background:var(--blue);color:white;border-radius:12px;margin-bottom:16px"><div style="font-size:20px;font-weight:800">'+(h2.name||'')+'</div><div style="font-size:13px;margin-top:4px">'+(h2.title||'')+'</div><div style="font-size:11px;opacity:.8;margin-top:4px">'+(h2.phone||'')+' | '+(h2.email||'')+' | '+(h2.location||'')+'</div></div>';
    if(d.cv.summary)html+=colSec('Profile Summary',d.cv.summary,'👤');
    if(d.cv.experience&&d.cv.experience.length){var exp='';d.cv.experience.forEach(function(e){exp+='<div style="margin-bottom:16px;padding:12px;background:var(--light);border-radius:8px;direction:ltr"><div style="font-weight:700;color:var(--blue)">'+e.title+' — '+e.company+'</div><div style="font-size:12px;color:var(--gray)">'+e.dates+'</div><ul style="margin-top:8px;padding-right:16px">'+(e.bullets||[]).map(function(b){return'<li style="font-size:12px;margin-bottom:4px">'+b+'</li>';}).join('')+'</ul></div>';});html+=colSec('الخبرة المهنية',exp,'💼');}
    if(d.cv.alignment_statement)html+=colSec('Alignment Statement',d.cv.alignment_statement,'🎯');
    document.getElementById('cvOut').innerHTML=html;
  }
  if(d.cover_letter){
    var cl=d.cover_letter;
    var html2='<div style="padding:16px;background:var(--blue);color:white;border-radius:12px;margin-bottom:16px"><div style="font-size:14px;font-weight:700">'+(cl.subject||'')+'</div></div><div style="direction:ltr;font-size:13px;line-height:1.8"><p style="margin-bottom:12px">Dear Hiring Manager,</p><p style="margin-bottom:12px">'+(cl.opening||'')+'</p>';
    if(cl.value_props&&cl.value_props.length){html2+='<div style="background:var(--blue-light);border-radius:8px;padding:12px;margin-bottom:12px"><strong style="color:var(--blue)">What I Bring:</strong><ul style="margin-top:8px;padding-right:20px">'+cl.value_props.map(function(v){return'<li style="margin-bottom:4px">'+v.text+'</li>';}).join('')+'</ul></div>';}
    (cl.body_paragraphs||[]).forEach(function(p2){html2+='<p style="margin-bottom:12px">'+p2+'</p>';});
    html2+='<p>'+(cl.closing||'')+'</p><p style="margin-top:16px;font-weight:700;color:var(--blue)">'+((cl.signature||{}).name||'')+'</p></div>';
    document.getElementById('clOut').innerHTML=html2;
  }
  if(d.linkedin){
    var li=d.linkedin,html3='';
    if(li.headline)html3+=colSec('Recommended Headline','<div style="direction:ltr;background:var(--blue-light);padding:12px;border-radius:8px">'+li.headline+'</div>','📝');
    if(li.about_section)html3+=colSec('About Section','<pre style="direction:ltr;font-size:12px;white-space:pre-wrap;line-height:1.6">'+li.about_section+'</pre>','👤');
    (li.messages||[]).forEach(function(m){html3+=colSec(m.type,'<div style="direction:ltr;font-size:13px;white-space:pre-wrap;line-height:1.6">'+(m.subject?'<strong>Subject:</strong> '+m.subject+'<br><br>':'')+m.body+'</div>','💬');});
    document.getElementById('liOut').innerHTML=html3;
  }
  if(d.star&&d.star.stories){
    var html4='';
    d.star.stories.forEach(function(s2,i2){html4+='<div class="stc"><div class="stch">⭐ قصة '+(i2+1)+': '+(s2.competency||'')+' — '+(s2.title||'')+'</div><div class="str"><div class="stl">SITUATION</div><div class="stt">'+(s2.situation||'')+'</div></div><div class="str"><div class="stl">TASK</div><div class="stt">'+(s2.task||'')+'</div></div><div class="str"><div class="stl">ACTION</div><div class="stt">'+(s2.action||'')+'</div></div><div class="str str-r"><div class="stl" style="color:var(--green)">RESULT</div><div class="stt" style="font-weight:600">'+(s2.result||'')+'</div></div><div style="padding:10px 12px;font-size:12px;color:var(--blue);font-style:italic;direction:ltr">💡 '+(s2.key_message||'')+'</div></div>';});
    document.getElementById('stOut').innerHTML=html4;
  }
  if(d.qa){
    var html5='';
    (d.qa.questions||[]).forEach(function(q2,i3){html5+='<div style="margin-bottom:20px;border:1px solid var(--border);border-radius:12px;overflow:hidden"><div style="background:var(--blue);color:white;padding:12px 16px;font-size:13px;font-weight:700">❓ سؤال '+(i3+1)+': '+(q2.question||'')+'</div><div style="padding:14px;font-size:13px;line-height:1.7;direction:ltr;background:var(--light)">'+(q2.answer||'')+'</div><div style="padding:10px 14px;font-size:12px;color:var(--blue);background:var(--blue-light)">💡 '+(q2.tip||'')+'</div></div>';});
    if(d.qa.questions_to_ask&&d.qa.questions_to_ask.length){html5+='<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:20px 0 10px">❓ أسئلة تطرحها أنت</h3>';d.qa.questions_to_ask.forEach(function(q3){html5+='<div style="padding:12px;border-right:4px solid var(--blue);background:var(--blue-light);border-radius:0 8px 8px 0;margin-bottom:8px;font-size:13px;direction:ltr">'+q3+'</div>';});}
    document.getElementById('qaOut').innerHTML=html5;
  }
  if(d.plan){
    var html6='';
    (d.plan.weeks||[]).forEach(function(w2){
      html6+='<h3 style="font-size:14px;font-weight:700;color:var(--blue);margin:16px 0 6px">'+(w2.week||'')+'</h3><p style="font-size:12px;color:var(--gray);margin-bottom:10px">'+(w2.focus||'')+'</p>';
      (w2.tasks||[]).forEach(function(t){var bg=t.priority==='HIGH'?'#FEE2E2':'var(--light)';html6+='<div style="display:flex;gap:8px;padding:10px;border-radius:8px;background:'+bg+';margin-bottom:6px;direction:ltr;align-items:flex-start"><span style="font-weight:700;font-size:11px;color:var(--blue);min-width:60px">'+(t.day||'')+'</span><span style="font-size:12px;flex:1">'+(t.action||'')+'</span>'+(t.platform?'<span style="font-size:10px;background:white;border-radius:6px;padding:2px 6px;color:var(--gray)">'+t.platform+'</span>':'')+'</div>';});
    });
    if(d.plan.pro_tips&&d.plan.pro_tips.length){html6+='<h3 style="font-size:14px;font-weight:700;color:var(--green);margin:20px 0 8px">💡 نصائح</h3>';d.plan.pro_tips.forEach(function(t2){html6+='<div style="padding:10px;border-right:4px solid var(--green);background:var(--green-light);border-radius:0 8px 8px 0;margin-bottom:6px;font-size:13px;direction:ltr">'+t2+'</div>';});}
    document.getElementById('plOut').innerHTML=html6;
  }
}

async function dlDoc(type){
  if(!gd||Object.keys(gd).length===0){alert('يرجى توليد الوثائق أولاً');return;}
  var body={},filename='';
  if(type==='cv'){body={cv_data:gd.cv};filename='CV_Optimized.docx';}
  else if(type==='cover_letter'){body={cl_data:gd.cover_letter,candidate_info:{name:gd.cName||''}};filename='Cover_Letter.docx';}
  else if(type==='strategy'){body={linkedin_data:gd.linkedin,star_data:gd.star,qa_data:gd.qa,plan_data:gd.plan};filename='Job_Strategy_Pack.docx';}
  try{
    var res=await fetch('/api/download/'+type,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});
    if(!res.ok)throw new Error('فشل التحميل');
    var blob=await res.blob();
    var url=URL.createObjectURL(blob);
    var a=document.createElement('a');a.href=url;a.download=filename;
    document.body.appendChild(a);a.click();document.body.removeChild(a);URL.revokeObjectURL(url);
  }catch(e){alert('خطأ: '+e.message);}
}
</script>
</body>
</html>"""

if __name__ == '__main__':
    print("""
╔══════════════════════════════════════════════════╗
║   CV ANALYZER PRO — النسخة المُصلَحة             ║
║   AI calls من المتصفح مباشرة ✅                  ║
║   http://localhost:5000                         ║
╚══════════════════════════════════════════════════╝
""")
    app.run(host='0.0.0.0', port=5000, debug=False)
